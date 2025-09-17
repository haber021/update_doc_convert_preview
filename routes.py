import os
import uuid
import re
from datetime import datetime, timedelta
from flask import render_template, request, redirect, url_for, flash, send_file, jsonify, abort, after_this_request
from flask import make_response
import mimetypes
import subprocess
import tempfile
import shutil
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename
from sqlalchemy import func, desc, or_
import pandas as pd
from app import app
from extensions import db
from models import User, Document, DocumentType, Role, AccessLog, SystemSettings, OwnerDocumentDetails
from utils import generate_qr_code, allowed_file, log_user_action, get_document_stats, predict_document_requests, resolve_document_file_path, get_document_folder
# Use the centralized email sender (Flask-Mail with SMTP fallback)
from email_sender import send_email
from flask import current_app
from sqlalchemy.exc import OperationalError
from sqlalchemy import text
from werkzeug.security import generate_password_hash
import secrets
import hmac
import hashlib
import base64
import string


def _has_temp_edit_grant(user_id: int, doc_id: int, minutes: int = 30) -> bool:
    """Return True if the user has a temporary edit grant for the document within the last 'minutes'.
    Grants are recorded as AccessLog entries with action='edit_grant'. Revocations are respected.
    """
    try:
        cutoff = datetime.utcnow() - timedelta(minutes=minutes)
        last_grant = AccessLog.query.filter(
            AccessLog.user_id == user_id,
            AccessLog.document_id == doc_id,
            AccessLog.action == 'edit_grant',
            AccessLog.timestamp >= cutoff
        ).order_by(desc(AccessLog.timestamp)).first()
        if not last_grant:
            return False
        # If a revoke occurred after the grant, it's no longer valid
        last_revoke = AccessLog.query.filter(
            AccessLog.user_id == user_id,
            AccessLog.document_id == doc_id,
            AccessLog.action == 'edit_revoke',
            AccessLog.timestamp >= last_grant.timestamp
        ).order_by(desc(AccessLog.timestamp)).first()
        return last_revoke is None
    except Exception:
        return False


def _find_soffice() -> str | None:
    """Try to find the LibreOffice/soffice executable across platforms.
    Honors app.config['SOFFICE_PATH'] if set.
    Returns absolute path or None if not found.
    """
    # 1) explicit config
    cfg = app.config.get('SOFFICE_PATH')
    if cfg and os.path.exists(cfg):
        return cfg
    # 2) PATH lookup
    path = shutil.which('soffice') or shutil.which('libreoffice')
    if path:
        return path
    # 3) Common Windows locations
    win_candidates = [
        r"C:\\Program Files\\LibreOffice\\program\\soffice.exe",
        r"C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
    ]
    for p in win_candidates:
        if os.path.exists(p):
            return p
    # 4) macOS
    mac_candidates = [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/LibreOffice",
    ]
    for p in mac_candidates:
        if os.path.exists(p):
            return p
    # 5) Linux common
    linux_candidates = [
        "/usr/bin/soffice",
        "/usr/local/bin/soffice",
        "/snap/bin/libreoffice",
    ]
    for p in linux_candidates:
        if os.path.exists(p):
            return p
    return None


def _b64url(data: bytes) -> str:
    return base64.urlsafe_b64encode(data).rstrip(b'=').decode('ascii')


def _unb64url(data: str) -> bytes:
    pad = '=' * (-len(data) % 4)
    return base64.urlsafe_b64decode((data + pad).encode('ascii'))


def _make_public_token(doc_id: int, expires_ts: int) -> str:
    """Create a short-lived signed token for anonymous file access.
    Format: b64url("doc_id:expires") + "." + b64url(HMACSHA256(payload, secret_key))
    """
    payload = f"{doc_id}:{expires_ts}".encode('utf-8')
    key = (app.secret_key or 'dev').encode('utf-8')
    sig = hmac.new(key, payload, hashlib.sha256).digest()
    return _b64url(payload) + '.' + _b64url(sig)


def _verify_public_token(token: str) -> int | None:
    try:
        part = token.split('.')
        if len(part) != 2:
            return None
        payload_b = _unb64url(part[0])
        sig_b = _unb64url(part[1])
        key = (app.secret_key or 'dev').encode('utf-8')
        good = hmac.new(key, payload_b, hashlib.sha256).digest()
        if not hmac.compare_digest(sig_b, good):
            return None
        payload = payload_b.decode('utf-8')
        doc_str, exp_str = payload.split(':', 1)
        if int(exp_str) < int(datetime.utcnow().timestamp()):
            return None
        return int(doc_str)
    except Exception:
        return None


def _convert_word_to_docx(src_path: str) -> str | None:
    """Convert legacy .doc/.odt to .docx using LibreOffice; fallback to Word COM on Windows.
    Returns absolute path to the generated .docx or None.
    """
    try:
        src_abs = os.path.abspath(src_path)
        if not os.path.exists(src_abs):
            return None
        # Try LibreOffice first
        soffice = _find_soffice()
        if soffice:
            outdir = tempfile.mkdtemp()
            lo_profile = tempfile.mkdtemp()
            lo_profile_url = 'file:///' + lo_profile.replace('\\', '/').replace(' ', '%20')
            def run_lo(convert_filter: str) -> subprocess.CompletedProcess:
                cmd = [
                    soffice,
                    '--headless','--nologo','--nodefault','--nofirststartwizard',
                    f"--env:UserInstallation={lo_profile_url}",
                    '--convert-to', convert_filter,
                    '--outdir', outdir,
                    src_abs,
                ]
                app.logger.info(f"Running LO doc->docx: {' '.join(cmd)}")
                return subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=180)
            # Try explicit filter name first, then generic
            proc = run_lo('docx:"MS Word 2007 XML"')
            if proc.returncode != 0:
                app.logger.warning(f"LO docx filter failed (code {proc.returncode}). stderr={proc.stderr.decode(errors='ignore')}")
                proc = run_lo('docx')
            if proc.returncode == 0:
                base = os.path.splitext(os.path.basename(src_abs))[0]
                docx_path = os.path.join(outdir, base + '.docx')
                if not os.path.exists(docx_path):
                    for name in os.listdir(outdir):
                        if name.lower().endswith('.docx'):
                            docx_path = os.path.join(outdir, name)
                            break
                if os.path.exists(docx_path):
                    return docx_path
        # Fallback to MS Word COM on Windows
        return _convert_word_to_docx_via_wordcom(src_abs)
    except Exception as e:
        app.logger.error(f'_convert_word_to_docx error: {e}')
        return None


def _convert_word_to_docx_via_wordcom(src_abs: str) -> str | None:
    """Windows-only: use Word COM to save as DOCX. Requires pywin32 and MS Word."""
    try:
        import sys
        if not sys.platform.startswith('win'):
            return None
        try:
            import win32com.client  # type: ignore
            import pythoncom  # type: ignore
        except Exception as e:
            app.logger.warning(f'pywin32 not available for Word COM docx conversion: {e}')
            return None
        out_dir = tempfile.mkdtemp()
        base = os.path.splitext(os.path.basename(src_abs))[0]
        out_path = os.path.join(out_dir, base + '.docx')
        try:
            pythoncom.CoInitialize()
            word = win32com.client.DispatchEx('Word.Application')
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(src_abs, ReadOnly=True)
            # wdFormatXMLDocument = 12 or 16? In modern Word: 12 is docx for SaveAs; 16 used by ExportAsFixedFormat PDF
            # Use 12 for SaveAs2 to DOCX
            doc.SaveAs2(out_path, FileFormat=12)
            doc.Close(False)
            word.Quit()
            pythoncom.CoUninitialize()
            if os.path.exists(out_path):
                app.logger.info('Converted DOC->DOCX via Word COM')
                return out_path
            return None
        except Exception as e:
            try:
                doc.Close(False)
            except Exception:
                pass
            try:
                word.Quit()
            except Exception:
                pass
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
            app.logger.error(f'Word COM doc->docx failed: {e}')
            return None
    except Exception as e:
        app.logger.error(f'_convert_word_to_docx_via_wordcom internal error: {e}')
        return None
        payload_b = _unb64url(part[0])
        sig_b = _unb64url(part[1])
        key = (app.secret_key or 'dev').encode('utf-8')
        good = hmac.new(key, payload_b, hashlib.sha256).digest()
        if not hmac.compare_digest(sig_b, good):
            return None
        payload = payload_b.decode('utf-8')
        doc_str, exp_str = payload.split(':', 1)
        if int(exp_str) < int(datetime.utcnow().timestamp()):
            return None
        return int(doc_str)
    except Exception:
        return None


def _convert_word_to_pdf(src_path: str) -> str | None:
    """Convert a .doc/.docx/.odt file to PDF using LibreOffice.
    Returns the absolute path to the generated PDF or None on failure.
    """
    try:
        soffice = _find_soffice()
        if not soffice:
            app.logger.warning('LibreOffice not found; cannot convert to PDF on-the-fly')
            return None
        src_abs = os.path.abspath(src_path)
        if not os.path.exists(src_abs):
            app.logger.error(f"Source file for conversion not found: {src_abs}")
            return None
        tmpdir = tempfile.mkdtemp()
        try:
            # Create a fresh temporary user profile to avoid first-run and lock issues
            lo_profile = tempfile.mkdtemp()
            lo_profile_url = 'file:///' + lo_profile.replace('\\', '/').replace(' ', '%20')

            def run_lo(convert_filter: str) -> subprocess.CompletedProcess:
                cmd = [
                    soffice,
                    '--headless', '--nologo', '--nodefault', '--nofirststartwizard',
                    f"--env:UserInstallation={lo_profile_url}",
                    '--convert-to', convert_filter,
                    '--outdir', tmpdir,
                    src_abs,
                ]
                app.logger.info(f"Running LibreOffice conversion: {' '.join(cmd)}")
                return subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=180)

            # First attempt with writer_pdf_Export
            proc = run_lo('pdf:writer_pdf_Export')
            if proc.returncode != 0:
                app.logger.warning(f"LibreOffice conversion via writer_pdf_Export failed (code {proc.returncode}). stderr={proc.stderr.decode(errors='ignore')}")
                # Fallback to generic 'pdf' filter
                proc = run_lo('pdf')
                if proc.returncode != 0:
                    app.logger.error(f"LibreOffice conversion failed (fallback) (code {proc.returncode}). stderr={proc.stderr.decode(errors='ignore')} stdout={proc.stdout.decode(errors='ignore')}")
                    # Try Windows MS Word COM automation as a last resort
                    pdf_from_com = _convert_word_to_pdf_via_wordcom(src_abs)
                    if pdf_from_com:
                        return pdf_from_com
                    return None
            base = os.path.splitext(os.path.basename(src_abs))[0]
            pdf_path = os.path.join(tmpdir, base + '.pdf')
            if not os.path.exists(pdf_path):
                # Sometimes LO writes uppercase/lowercase differently; search for any .pdf in tmpdir
                for name in os.listdir(tmpdir):
                    if name.lower().endswith('.pdf'):
                        pdf_path = os.path.join(tmpdir, name)
                        break
            if not os.path.exists(pdf_path):
                app.logger.error('Conversion reported success but PDF not found')
                return None
            return pdf_path
        except subprocess.TimeoutExpired:
            app.logger.error('LibreOffice conversion timed out')
            return None
        except Exception as e:
            app.logger.error(f'LibreOffice conversion error: {e}')
            return None
    except Exception as e:
        app.logger.error(f'_convert_word_to_pdf internal error: {e}')
        return None


def _convert_word_to_pdf_via_wordcom(src_abs: str) -> str | None:
    """Windows-only fallback using Microsoft Word COM to export to PDF.
    Requires Microsoft Word installed and pywin32 package.
    Returns path to generated PDF or None.
    """
    try:
        import sys
        if not sys.platform.startswith('win'):
            return None
        try:
            import win32com.client  # type: ignore
            import pythoncom  # type: ignore
        except Exception as e:
            app.logger.warning(f'pywin32 not available for Word COM conversion: {e}')
            return None
        if not os.path.exists(src_abs):
            return None
        out_dir = tempfile.mkdtemp()
        base = os.path.splitext(os.path.basename(src_abs))[0]
        pdf_path = os.path.join(out_dir, base + '.pdf')
        try:
            pythoncom.CoInitialize()
            word = win32com.client.DispatchEx('Word.Application')
            word.Visible = False
            word.DisplayAlerts = 0  # wdAlertsNone
            doc = word.Documents.Open(src_abs, ReadOnly=True)
            # wdExportFormatPDF = 17
            # wdExportOptimizeForPrint = 0
            # OpenAfterExport=False, OptimizeFor=0, Range=0 (wdExportAllDocument), From=1, To=1, Item=0, IncludeDocProps=True, KeepIRM=True, CreateBookmarks=1, DocStructureTags=True, BitmapMissingFonts=True, UseISO19005_1=False
            doc.ExportAsFixedFormat(pdf_path, 17, False, 0, 0, 1, 1, 0, True, True, 1, True, False, False, False)
            doc.Close(False)
            word.Quit()
            pythoncom.CoUninitialize()
            if os.path.exists(pdf_path):
                app.logger.info('Converted via Microsoft Word COM automation')
                return pdf_path
            app.logger.error('Word COM reported success but PDF not found')
            return None
        except Exception as e:
            try:
                doc.Close(False)
            except Exception:
                pass
            try:
                word.Quit()
            except Exception:
                pass
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass
            app.logger.error(f'Word COM conversion failed: {e}')
            return None
    except Exception as e:
        app.logger.error(f'_convert_word_to_pdf_via_wordcom internal error: {e}')
        return None


def _has_temp_grant(user_id: int, doc_id: int, kind: str, minutes: int = 30) -> bool:
    """Check if user has a temporary grant for a specific kind (download, print, update, open, email, owner_edit).
    Respects explicit revocations ({kind}_revoke) after the grant timestamp.
    """
    try:
        cutoff = datetime.utcnow() - timedelta(minutes=minutes)
        action = f"{kind}_grant"
        last_grant = AccessLog.query.filter(
            AccessLog.user_id == user_id,
            AccessLog.document_id == doc_id,
            AccessLog.action == action,
            AccessLog.timestamp >= cutoff
        ).order_by(desc(AccessLog.timestamp)).first()
        if not last_grant:
            return False
        revoke_action = f"{kind}_revoke"
        revoked = AccessLog.query.filter(
            AccessLog.user_id == user_id,
            AccessLog.document_id == doc_id,
            AccessLog.action == revoke_action,
            AccessLog.timestamp >= last_grant.timestamp
        ).order_by(desc(AccessLog.timestamp)).first()
        return revoked is None
    except Exception:
        return False

@app.route('/')
@login_required
def dashboard():
    # Get dashboard statistics
    try:
        total_documents = Document.query.filter_by(is_active=True).count()
    except OperationalError:
        # Attempt light migration: add document_number and backfill, then retry
        try:
            conn = db.engine.connect()
            res = conn.execute(text("PRAGMA table_info('document')"))
            cols = [row[1] for row in res.fetchall()]
            if 'document_number' not in cols:
                current_app.logger.info("Adding 'document_number' column (runtime migration)")
                conn.execute(text("ALTER TABLE document ADD COLUMN document_number VARCHAR"))
                docs = conn.execute(text("SELECT id FROM document")).fetchall()
                for (doc_id,) in docs:
                    code = f"DOC{int(doc_id):06d}"
                    conn.execute(text("UPDATE document SET document_number = :code WHERE id = :id"), {'code': code, 'id': doc_id})
            conn.close()
        except Exception as e:
            current_app.logger.warning(f"Runtime migration failed: {e}")
        # retry once
        total_documents = Document.query.filter_by(is_active=True).count()
    active_users_today = User.query.join(AccessLog).filter(
        AccessLog.timestamp >= datetime.utcnow().date()
    ).distinct().count()
    
    documents_accessed_today = AccessLog.query.filter(
        AccessLog.timestamp >= datetime.utcnow().date(),
        AccessLog.action.in_(['view', 'download'])
    ).count()
    
    # Get recent documents for current user
    if current_user.has_role('Student'):
        recent_docs = current_user.owned_documents.filter_by(is_active=True).order_by(desc(Document.created_at)).limit(5).all()
    else:
        recent_docs = Document.query.filter_by(is_active=True).order_by(desc(Document.created_at)).limit(5).all()
    
    # Get predictive insights
    predictions = predict_document_requests()
    
    return render_template('dashboard.html',
                         total_documents=total_documents,
                         active_users_today=active_users_today,
                         documents_accessed_today=documents_accessed_today,
                         recent_documents=recent_docs,
                         predictions=predictions)

@app.route('/documents')
@login_required
def documents():
    page = request.args.get('page', 1, type=int)
    search = request.args.get('search', '')
    doc_type_filter = request.args.get('type', '')
    status_filter = request.args.get('status', 'active')
    
    query = Document.query
    
    # Apply role-based filtering
    if current_user.has_role('Student'):
        query = query.filter(Document.owner_id == current_user.id)
    
    # Apply filters
    if search:
        # perform case-insensitive-ish search against title, filename, document_number, and owner name/ID
        q = f"%{search}%"
        # join owner for searching owner names
        try:
            query = query.outerjoin(User, Document.owner)
        except Exception:
            # fallback if relationship name differs
            pass
        # Join owner document details for searching custom names
        try:
            query = query.outerjoin(OwnerDocumentDetails, OwnerDocumentDetails.document_id == Document.id)
        except Exception:
            pass
        query = query.filter(or_(
            Document.title.ilike(q),
            Document.filename.ilike(q),
            Document.document_number.ilike(q),
            User.first_name.ilike(q),
            User.last_name.ilike(q),
            OwnerDocumentDetails.first_name.ilike(q),
            OwnerDocumentDetails.middle_name.ilike(q),
            OwnerDocumentDetails.last_name.ilike(q),
            # student_id may be numeric; cast for safety
            func.cast(User.student_id, db.String).ilike(q)
        ))
    
    if doc_type_filter:
        query = query.filter(Document.document_type_id == doc_type_filter)
    
    if status_filter == 'active':
        query = query.filter(Document.is_active == True)
    elif status_filter == 'inactive':
        query = query.filter(Document.is_active == False)
    
    # Pagination
    documents = query.order_by(desc(Document.created_at)).paginate(
        page=page, per_page=10, error_out=False
    )
    
    # Get document types for filter dropdown
    document_types = DocumentType.query.all()
    
    return render_template('documents.html',
                         documents=documents,
                         document_types=document_types,
                         search=search,
                         doc_type_filter=doc_type_filter,
                         status_filter=status_filter)

@app.route('/upload', methods=['GET', 'POST'])
@login_required
def upload_document():
    if not current_user.can_upload_documents():
        flash('You do not have permission to upload documents.', 'error')
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        title = request.form.get('title')
        # If no student/employee ID provided, default to the logged-in user's ID
        student_id = request.form.get('student_id') or current_user.student_id
        doc_type_id = request.form.get('document_type')
        file = request.files.get('file')

        # Validation
        if not all([title, doc_type_id, file]):
            flash('Please fill in all required fields.', 'error')
            return redirect(request.url)

        if not allowed_file(file.filename):
            flash('Invalid file type. Please upload PDF, DOC, DOCX, or image files.', 'error')
            return redirect(request.url)

        # Find document owner: default to the uploader (current_user). If a student_id
        # was supplied and it differs from current_user.student_id, try to resolve that to a user.
        owner = current_user
        if student_id and str(student_id) != str(current_user.student_id):
            student = User.query.filter_by(student_id=student_id).first()
            if student:
                owner = student
            else:
                flash('Student/Employee ID not found.', 'error')
                return redirect(request.url)

        # Save file
        filename = secure_filename(file.filename)
        extension = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
        folder_name = get_document_folder(extension)
        folder = os.path.join(app.config['UPLOAD_FOLDER'], folder_name)
        os.makedirs(folder, exist_ok=True)
        unique_filename = f"{uuid.uuid4()}_{filename}"
        file_path = os.path.join(folder, unique_filename)
        file.save(file_path)

        # Create document record
        document = Document(
            title=title,
            filename=filename,
            file_path=file_path,
            file_size=os.path.getsize(file_path),
            file_type=file.content_type,
            document_type_id=doc_type_id,
            uploaded_by_id=current_user.id,
            owner_id=owner.id
        )

        db.session.add(document)
        db.session.flush()  # To get the document ID

        # Auto-convert legacy .doc to .docx on upload for better preview compatibility
        try:
            orig_name = (document.filename or '')
            ext = os.path.splitext(orig_name)[1].lower()
            if ext == '.doc':
                src_for_conv = document.file_path or locals().get('file_path') or ''
                docx_tmp = _convert_word_to_docx(src_for_conv)
                if docx_tmp and os.path.exists(docx_tmp):
                    # Persist the converted file into uploads/converted/ with a safe unique name
                    dest_dir = os.path.join(app.config.get('UPLOAD_FOLDER', 'uploads'), 'converted')
                    os.makedirs(dest_dir, exist_ok=True)
                    safe_base = secure_filename(os.path.splitext(orig_name)[0]) or f'doc_{document.id}'
                    new_name = f"{safe_base}.docx"
                    dest_path = os.path.join(dest_dir, new_name)
                    # Avoid collision by appending a short suffix
                    if os.path.exists(dest_path):
                        new_name = f"{safe_base}_{uuid.uuid4().hex[:8]}.docx"
                        dest_path = os.path.join(dest_dir, new_name)
                    shutil.move(docx_tmp, dest_path)

                    # Update Document to point to the new DOCX
                    document.file_path = dest_path
                    document.filename = new_name
                    document.file_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

                    # Optionally remove original .doc from disk if it exists
                    try:
                        if src_for_conv and os.path.exists(src_for_conv):
                            os.remove(src_for_conv)
                    except Exception:
                        pass
        except Exception as e:
            app.logger.warning(f'DOC upload auto-convert to DOCX failed for doc {document.id}: {e}')

        # Save owner document details if provided
        try:
            ofn = (request.form.get('owner_first_name') or '').strip()
            omn = (request.form.get('owner_middle_name') or '').strip()
            oln = (request.form.get('owner_last_name') or '').strip()
            oaddr = (request.form.get('owner_address') or '').strip()
            if any([ofn, omn, oln, oaddr]):
                odd = OwnerDocumentDetails(
                    document_id=document.id,
                    first_name=ofn or None,
                    middle_name=omn or None,
                    last_name=oln or None,
                    address=oaddr or None
                )
                db.session.add(odd)
        except Exception as e:
            app.logger.warning(f'Failed to record owner details for document {document.id}: {e}')

        # Assign a human-friendly document number (e.g., DOC000001)
        try:
            document.document_number = f"DOC{document.id:06d}"
        except Exception:
            document.document_number = str(document.id)

        # Generate QR code
        qr_path = generate_qr_code(document.id)
        document.qr_code_path = qr_path
        document.qr_code_data = f"{request.url_root}document/{document.id}"

        db.session.commit()

        # Log the action
        log_user_action(current_user.id, 'upload', document.id, f"Uploaded document: {title}")

        flash('Document uploaded successfully!', 'success')
        return redirect(url_for('documents'))
    
    # GET request - show upload form
    document_types = DocumentType.query.all()
    return render_template('upload.html', document_types=document_types)

@app.route('/document/<int:doc_id>')
@login_required
def document_detail(doc_id):
    document = Document.query.get_or_404(doc_id)
    days_since_created = (datetime.utcnow() - document.created_at).days
    # Fetch related documents server-side to avoid using model classes in Jinja templates
    related_docs = Document.query.filter_by(
        owner_id=document.owner_id,
        is_active=True,
        document_type_id=document.document_type_id
    ).filter(Document.id != document.id).order_by(desc(Document.created_at)).limit(5).all()

    # Determine preview kind safely (image/pdf/word/other)
    preview_kind = 'other'
    # Prefer stored file_type if available
    ft = (document.file_type or '').lower()
    if ft:
        if ft.startswith('image/'):
            preview_kind = 'image'
        elif 'pdf' in ft:
            preview_kind = 'pdf'
        elif 'word' in ft or 'msword' in ft or 'officedocument' in ft:
            preview_kind = 'word'
    else:
        # Fallback to filename extension
        ext = os.path.splitext(document.filename or '')[1].lower()
        if ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp']:
            preview_kind = 'image'
        elif ext == '.pdf':
            preview_kind = 'pdf'
        elif ext in ['.doc', '.docx', '.odt']:
            preview_kind = 'word'

    # Compute edit permission: admin or temporary grant
    can_edit = current_user.can_access_admin() or _has_temp_edit_grant(current_user.id, document.id)

    # Gather pending edit requests for admins
    pending_requests = []
    if current_user.can_access_admin():
        try:
            pending_requests = AccessLog.query.filter_by(document_id=document.id, action='request_edit').order_by(desc(AccessLog.timestamp)).limit(20).all()
        except Exception:
            pending_requests = []

    # Compute edit expiration for non-admin users with a grant (30 minutes window)
    edit_expiration = None
    if can_edit and not current_user.can_access_admin():
        try:
            last_grant = AccessLog.query.filter_by(user_id=current_user.id, document_id=document.id, action='edit_grant').order_by(desc(AccessLog.timestamp)).first()
            if last_grant:
                exp = last_grant.timestamp + timedelta(minutes=30)
                if exp > datetime.utcnow():
                    edit_expiration = exp
        except Exception:
            edit_expiration = None

    # Per-action temporary grants
    try:
        has_download_grant = current_user.can_access_admin() or _has_temp_grant(current_user.id, document.id, 'download')
    except Exception:
        has_download_grant = current_user.can_access_admin()
    try:
        has_print_grant = current_user.can_access_admin() or _has_temp_grant(current_user.id, document.id, 'print')
    except Exception:
        has_print_grant = current_user.can_access_admin()
    try:
        has_update_grant = current_user.can_access_admin() or _has_temp_grant(current_user.id, document.id, 'update')
    except Exception:
        has_update_grant = current_user.can_access_admin()
    try:
        has_open_grant = current_user.can_access_admin() or _has_temp_grant(current_user.id, document.id, 'open')
    except Exception:
        has_open_grant = current_user.can_access_admin()
    try:
        has_email_grant = current_user.can_access_admin() or _has_temp_grant(current_user.id, document.id, 'email')
    except Exception:
        has_email_grant = current_user.can_access_admin()
    try:
        has_owner_edit_grant = current_user.can_access_admin() or _has_temp_grant(current_user.id, document.id, 'owner_edit')
    except Exception:
        has_owner_edit_grant = current_user.can_access_admin()

    # Compute expiration times for each grant (non-admins)
    download_expiration = None
    print_expiration = None
    update_expiration = None
    email_expiration = None
    open_expiration = None
    owner_edit_expiration = None
    if not current_user.can_access_admin():
        try:
            if has_download_grant:
                last_grant = AccessLog.query.filter_by(user_id=current_user.id, document_id=document.id, action='download_grant').order_by(desc(AccessLog.timestamp)).first()
                if last_grant:
                    exp = last_grant.timestamp + timedelta(minutes=30)
                    if exp > datetime.utcnow():
                        download_expiration = exp
        except Exception:
            download_expiration = None
        try:
            if has_print_grant:
                last_grant = AccessLog.query.filter_by(user_id=current_user.id, document_id=document.id, action='print_grant').order_by(desc(AccessLog.timestamp)).first()
                if last_grant:
                    exp = last_grant.timestamp + timedelta(minutes=30)
                    if exp > datetime.utcnow():
                        print_expiration = exp
        except Exception:
            print_expiration = None
        try:
            if has_update_grant:
                last_grant = AccessLog.query.filter_by(user_id=current_user.id, document_id=document.id, action='update_grant').order_by(desc(AccessLog.timestamp)).first()
                if last_grant:
                    exp = last_grant.timestamp + timedelta(minutes=30)
                    if exp > datetime.utcnow():
                        update_expiration = exp
        except Exception:
            update_expiration = None
        try:
            if has_email_grant:
                last_grant = AccessLog.query.filter_by(user_id=current_user.id, document_id=document.id, action='email_grant').order_by(desc(AccessLog.timestamp)).first()
                if last_grant:
                    exp = last_grant.timestamp + timedelta(minutes=30)
                    if exp > datetime.utcnow():
                        email_expiration = exp
        except Exception:
            email_expiration = None
        try:
            if has_open_grant:
                last_grant = AccessLog.query.filter_by(user_id=current_user.id, document_id=document.id, action='open_grant').order_by(desc(AccessLog.timestamp)).first()
                if last_grant:
                    exp = last_grant.timestamp + timedelta(minutes=30)
                    if exp > datetime.utcnow():
                        open_expiration = exp
        except Exception:
            open_expiration = None
        try:
            if has_owner_edit_grant:
                last_grant = AccessLog.query.filter_by(user_id=current_user.id, document_id=document.id, action='owner_edit_grant').order_by(desc(AccessLog.timestamp)).first()
                if last_grant:
                    exp = last_grant.timestamp + timedelta(minutes=30)
                    if exp > datetime.utcnow():
                        owner_edit_expiration = exp
        except Exception:
            owner_edit_expiration = None

    # Determine if there are unresolved requests for specific actions by this user
    kinds = ['download', 'print', 'update', 'open', 'email', 'owner_edit', 'edit']
    unresolved_requests = {k: False for k in kinds}
    if not current_user.can_access_admin():
        for k in kinds:
            try:
                if k == 'owner_edit':
                    req_action = 'request_owner_edit'
                    grant_action = 'owner_edit_grant'
                    reject_action = 'request_owner_edit_reject'
                elif k == 'edit':
                    req_action = 'request_edit'
                    grant_action = 'edit_grant'
                    reject_action = 'request_edit_reject'
                else:
                    req_action = f'request_{k}'
                    grant_action = f'{k}_grant'
                    reject_action = f'{req_action}_reject'

                last_req = AccessLog.query.filter_by(
                    user_id=current_user.id,
                    document_id=document.id,
                    action=req_action
                ).order_by(desc(AccessLog.timestamp)).first()

                if last_req:
                    resolution = AccessLog.query.filter(
                        AccessLog.user_id == current_user.id,
                        AccessLog.document_id == document.id,
                        AccessLog.action.in_([grant_action, reject_action]),
                        AccessLog.timestamp >= last_req.timestamp
                    ).first()
                    unresolved_requests[k] = (resolution is None)
            except Exception:
                unresolved_requests[k] = False

    # Generate a short-lived public URL for external viewers (e.g., Office/Google)
    try:
        expires_ts = int((datetime.utcnow() + timedelta(minutes=5)).timestamp())
        public_token = _make_public_token(document.id, expires_ts)
        public_view_url = url_for('public_document_view', token=public_token, _external=True)
    except Exception:
        public_view_url = None

    return render_template('document_detail.html',
                           document=document,
                           days_since_created=days_since_created,
                           related_docs=related_docs,
                           preview_kind=preview_kind,
                           can_edit=can_edit,
                           pending_requests=pending_requests,
                           edit_expiration=edit_expiration,
                           has_download_grant=has_download_grant,
                           has_print_grant=has_print_grant,
                           has_update_grant=has_update_grant,
                           has_open_grant=has_open_grant,
                           has_email_grant=has_email_grant,
                           download_expiration=download_expiration,
                           print_expiration=print_expiration,
                           update_expiration=update_expiration,
                           email_expiration=email_expiration,
                           open_expiration=open_expiration,
                           has_owner_edit_grant=has_owner_edit_grant,
                           owner_edit_expiration=owner_edit_expiration,
                           unresolved_requests=unresolved_requests,
                           public_view_url=public_view_url)

@app.route('/download/<int:doc_id>')
@login_required
def download_document(doc_id):
    document = Document.query.get_or_404(doc_id)

    # Check permissions
    if current_user.has_role('Student') and document.owner_id != current_user.id:
        abort(403)

    # Admins can always download. Others require a temporary grant.
    if not current_user.can_access_admin():
        if not _has_temp_grant(current_user.id, document.id, 'download'):
            flash('Download requires administrator approval. A request is needed.', 'warning')
            return redirect(url_for('document_detail', doc_id=document.id))

    # Resolve file path
    file_path = document.file_path or ''
    if not os.path.exists(file_path):
        alt = os.path.join(app.config.get('UPLOAD_FOLDER', 'uploads'), file_path)
        if os.path.exists(alt):
            file_path = alt
        else:
            candidate = os.path.join(app.config.get('UPLOAD_FOLDER', 'uploads'), document.filename or '')
            if os.path.exists(candidate):
                file_path = candidate
            else:
                app.logger.error(f"File not found for document {document.id}: {document.file_path}")
                abort(404)

    # Log the download
    log_user_action(current_user.id, 'download', doc_id, f"Downloaded document: {document.title}")
    document.increment_access()

    return send_file(file_path, as_attachment=True, download_name=document.filename)


@app.route('/document/<int:doc_id>/view')
@login_required
def view_document(doc_id):
    """Serve the document for inline viewing (no attachment) when possible."""
    document = Document.query.get_or_404(doc_id)

    # Permission check: students can only view their own documents
    if current_user.has_role('Student') and document.owner_id != current_user.id:
        abort(403)

    # Enforce 'open' grant for non-admins to view inline
    if not current_user.can_access_admin():
        if not _has_temp_grant(current_user.id, document.id, 'open'):
            flash('Opening requires administrator approval. A request is needed.', 'warning')
            return redirect(url_for('document_detail', doc_id=document.id))

    # Log the view and increment access count
    log_user_action(current_user.id, 'view', document.id, f"Viewed document: {document.title}")
    document.increment_access()

    # Serve file inline (browser will render PDFs/images if supported)
    try:
        # Resolve file path candidates
        file_path = document.file_path or ''
        if not os.path.exists(file_path):
            alt = os.path.join(app.config.get('UPLOAD_FOLDER', 'uploads'), file_path)
            if os.path.exists(alt):
                file_path = alt
            else:
                candidate = os.path.join(app.config.get('UPLOAD_FOLDER', 'uploads'), document.filename or '')
                if os.path.exists(candidate):
                    file_path = candidate
                else:
                    app.logger.error(f"File not found for document {document.id}: {document.file_path}")
                    abort(404)

        mimetype = (document.file_type or mimetypes.guess_type(file_path)[0] or '').lower()

        # Serve raw bytes if explicitly requested (used for client-side DOCX rendering)
        if request.args.get('raw') == '1':
            guessed_mime = mimetypes.guess_type(file_path)[0] or 'application/octet-stream'
            resp = make_response(send_file(file_path, as_attachment=False, mimetype=guessed_mime))
            resp.headers['Content-Disposition'] = f'inline; filename="{document.filename}"'
            resp.headers['X-Content-Type-Options'] = 'nosniff'
            return resp

        # If already PDF, serve inline
        if 'pdf' in mimetype or file_path.lower().endswith('.pdf'):
            resp = make_response(send_file(file_path, as_attachment=False, mimetype='application/pdf'))
            resp.headers['Content-Disposition'] = f'inline; filename="{document.filename or "document.pdf"}"'
            return resp

        # If Word doc, attempt conversion using LibreOffice and stream PDF inline
        if file_path.lower().endswith(('.doc', '.docx', '.odt')):
            pdf_path = _convert_word_to_pdf(file_path)
            if not pdf_path:
                # graceful fallback: show help with download link
                return f"""
<html>
<head><title>Preview not available</title></head>
<body>
<h3>Preview not available</h3>
<p>The document cannot be previewed inline at this time. Please <a href="{url_for('download_document', doc_id=document.id)}" target="_blank">download the file</a> to view it.</p>
</body>
</html>
""", 200, {'Content-Type': 'text/html'}

            @after_this_request
            def _cleanup_temp(response):
                try:
                    if pdf_path and os.path.exists(pdf_path):
                        os.remove(pdf_path)
                except Exception:
                    pass
                return response

            base = os.path.splitext(os.path.basename(file_path))[0]
            resp = make_response(send_file(pdf_path, as_attachment=False, mimetype='application/pdf'))
            resp.headers['Content-Disposition'] = f'inline; filename="{base}.pdf"'
            return resp

        # For images, serve inline
        if mimetype.startswith('image/'):
            resp = make_response(send_file(file_path, as_attachment=False, mimetype=mimetype))
            resp.headers['Content-Disposition'] = f'inline; filename="{document.filename}"'
            resp.headers['X-Content-Type-Options'] = 'nosniff'
            return resp

        # For other types, serve as inline (browser may download)
        guessed_mime = mimetypes.guess_type(file_path)[0] or 'application/octet-stream'
        resp = make_response(send_file(file_path, as_attachment=False, mimetype=guessed_mime))
        resp.headers['Content-Disposition'] = f'inline; filename="{document.filename}"'
        resp.headers['X-Content-Type-Options'] = 'nosniff'
        return resp
    except Exception as e:
        app.logger.error(f"Error serving file for view: {e}")
        abort(404)


@app.route('/document/<int:doc_id>/print')
@login_required
def print_document(doc_id):
    """Render a print-friendly page for the document and trigger the browser print dialog."""
    document = Document.query.get_or_404(doc_id)

    # Permission check: students can only print their own documents
    if current_user.has_role('Student') and document.owner_id != current_user.id:
        abort(403)

    # Admins can always print; others require a temporary print grant
    if not current_user.can_access_admin():
        if not _has_temp_grant(current_user.id, document.id, 'print'):
            flash('Printing requires administrator approval. A request is needed.', 'warning')
            return redirect(url_for('document_detail', doc_id=document.id))

    # Determine preview kind (reuse same logic as detail)
    preview_kind = 'other'
    ft = (document.file_type or '').lower()
    if ft:
        if ft.startswith('image/'):
            preview_kind = 'image'
        elif 'pdf' in ft:
            preview_kind = 'pdf'
        elif 'word' in ft or 'msword' in ft or 'officedocument' in ft:
            preview_kind = 'word'
    else:
        ext = os.path.splitext(document.filename or '')[1].lower()
        if ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp']:
            preview_kind = 'image'
        elif ext == '.pdf':
            preview_kind = 'pdf'
        elif ext in ['.doc', '.docx', '.odt']:
            preview_kind = 'word'

    # Choose embedding strategy so browsers can render content inline
    # For images, embed the view endpoint (serves image inline)
    if preview_kind == 'image':
        file_url = url_for('view_document', doc_id=document.id) + '?_ts=' + str(int(datetime.utcnow().timestamp()))
        return render_template('print_document.html', document=document, preview_kind='image', file_url=file_url)

    # For PDFs, use the PDF endpoint (serves inline)
    if preview_kind == 'pdf':
        file_url = url_for('document_pdf', doc_id=document.id) + '?_ts=' + str(int(datetime.utcnow().timestamp()))
        return render_template('print_document.html', document=document, preview_kind='pdf', file_url=file_url)

    # For Word documents, attempt to use LibreOffice conversion; if not available,
    # for .docx try extracting text and render HTML for printing.
    ext = os.path.splitext(document.filename or '')[1].lower()
    soffice = _find_soffice()
    if soffice:
        # Let document_pdf handle conversion and embedding
        file_url = url_for('document_pdf', doc_id=document.id) + '?_ts=' + str(int(datetime.utcnow().timestamp()))
        return render_template('print_document.html', document=document, preview_kind='pdf', file_url=file_url)

    # No soffice: if .docx, try to extract text and show a printable HTML page
    if ext == '.docx':
        try:
            from docx import Document as DocxDocument
            docx_path = document.file_path or ''
            if not os.path.exists(docx_path):
                alt = os.path.join(app.config.get('UPLOAD_FOLDER', 'uploads'), docx_path)
                if os.path.exists(alt):
                    docx_path = alt
                else:
                    candidate = os.path.join(app.config.get('UPLOAD_FOLDER', 'uploads'), document.filename or '')
                    if os.path.exists(candidate):
                        docx_path = candidate
            paragraphs = []
            if os.path.exists(docx_path):
                dd = DocxDocument(docx_path)
                for p in dd.paragraphs:
                    if p.text and p.text.strip():
                        paragraphs.append(p.text)
            return render_template('print_docx.html', document=document, paragraphs=paragraphs)
        except Exception as e:
            app.logger.warning(f"docx extraction failed: {e}")

    # Fallback: use PDF endpoint (which may redirect to download) or download directly
    file_url = url_for('document_pdf', doc_id=document.id) + '?_ts=' + str(int(datetime.utcnow().timestamp()))
    return render_template('print_document.html', document=document, preview_kind='pdf', file_url=file_url)


@app.route('/document/<int:doc_id>/review')
@login_required
def review_document(doc_id):
    document = Document.query.get_or_404(doc_id)

    # Permission check: students can only review their own documents
    if current_user.has_role('Student') and document.owner_id != current_user.id:
        abort(403)

    # Log review and increment access count
    try:
        log_user_action(current_user.id, 'review', document.id, f"Reviewed document: {document.title}")
        document.increment_access()
    except Exception:
        pass

    # Use the consolidated PDF endpoint for a consistent inline preview.
    # This will stream PDFs inline, convert Word/images to PDF when possible, or fallback to download.
    return redirect(url_for('document_pdf', doc_id=document.id))

@app.route('/document/<int:doc_id>/edit', methods=['GET'])
@login_required
def edit_document(doc_id):
    document = Document.query.get_or_404(doc_id)

    # Permissions: admin or temporary grant approved by admin
    if not (current_user.can_access_admin() or _has_temp_edit_grant(current_user.id, document.id)):
        abort(403)

    # Resolve file path
    file_path = resolve_document_file_path(document) or (document.file_path or '')
    if not os.path.exists(file_path):
        flash('File not found for editing.', 'error')
        return redirect(url_for('document_detail', doc_id=document.id))

    ext = os.path.splitext(document.filename or '')[1].lower()
    editor_kind = None
    payload = {}

    try:
        if ext in ['.txt']:
            editor_kind = 'text'
            # Limit text size to avoid heavy files
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            payload = {'text': content}
        elif ext in ['.csv']:
            editor_kind = 'table'
            df = pd.read_csv(file_path, dtype=str, keep_default_na=False, low_memory=False)
            columns = list(df.columns.astype(str))
            rows = df.values.tolist()
            payload = {'columns': columns, 'rows': rows, 'format': 'csv'}
        elif ext in ['.xlsx']:
            editor_kind = 'table'
            df = pd.read_excel(file_path, dtype=str)  # uses openpyxl engine by default if installed
            df = df.fillna('')
            columns = list(df.columns.astype(str))
            rows = df.values.tolist()
            payload = {'columns': columns, 'rows': rows, 'format': 'xlsx'}
        elif ext in ['.docx']:
            editor_kind = 'text'
            try:
                from docx import Document as DocxDocument
            except ImportError:
                flash('DOCX editing requires python-docx. Please install it on the server (pip install python-docx).', 'error')
                return redirect(url_for('document_detail', doc_id=document.id))
            try:
                dd = DocxDocument(file_path)
                paragraphs = [p.text for p in dd.paragraphs]
                payload = {'text': '\n\n'.join(paragraphs), 'format': 'docx'}
            except Exception as e:
                app.logger.error(f'DOCX open failed: {e}')
                flash('Failed to open the DOCX file for editing.', 'error')
                return redirect(url_for('document_detail', doc_id=document.id))
        else:
            flash('Editing is supported only for .txt, .csv, .xlsx, and .docx files at this time.', 'warning')
            return redirect(url_for('document_detail', doc_id=document.id))
    except Exception as e:
        app.logger.error(f"Failed to open file for editing: {e}")
        flash('Failed to open the file for editing.', 'error')
        return redirect(url_for('document_detail', doc_id=document.id))

    # Compute edit expiration for non-admin users
    edit_expires_at = None
    if not current_user.can_access_admin():
        try:
            last_grant = AccessLog.query.filter_by(user_id=current_user.id, document_id=document.id, action='edit_grant').order_by(desc(AccessLog.timestamp)).first()
            if last_grant:
                exp = last_grant.timestamp + timedelta(minutes=30)
                if exp > datetime.utcnow():
                    edit_expires_at = exp.strftime('%Y-%m-%dT%H:%M:%SZ')
        except Exception:
            pass

    return render_template('edit_document.html', document=document, editor_kind=editor_kind, payload=payload, edit_expires_at=edit_expires_at)


@app.route('/document/<int:doc_id>/edit/save', methods=['POST'])
@login_required
def save_document_edit(doc_id):
    document = Document.query.get_or_404(doc_id)

    if not (current_user.can_access_admin() or _has_temp_edit_grant(current_user.id, document.id)):
        return jsonify({'success': False, 'message': 'Access denied'}), 403

    file_path = resolve_document_file_path(document) or (document.file_path or '')
    if not os.path.exists(file_path):
        return jsonify({'success': False, 'message': 'File not found'}), 404

    data = request.get_json(silent=True) or {}
    kind = data.get('kind')

    # Create backup
    try:
        backup_path = file_path + '.bak'
        try:
            if os.path.exists(backup_path):
                os.remove(backup_path)
        except Exception:
            pass
        import shutil as _sh
        _sh.copyfile(file_path, backup_path)
    except Exception as e:
        app.logger.warning(f"Could not create backup before edit: {e}")

    try:
        if kind == 'text':
            fmt = data.get('format')
            text = data.get('text', '')
            if fmt == 'docx':
                try:
                    from docx import Document as DocxDocument
                except ImportError:
                    return jsonify({'success': False, 'message': 'DOCX editing requires python-docx on the server'}), 500
                try:
                    docx = DocxDocument()
                    # Write paragraphs; note: this will not preserve original formatting/images
                    for i, para in enumerate(text.split('\n\n')):
                        if i == 0 and docx.paragraphs:
                            docx.paragraphs[0].text = para
                        else:
                            docx.add_paragraph(para)
                    docx.save(file_path)
                except Exception as e:
                    app.logger.error(f'DOCX save failed: {e}')
                    return jsonify({'success': False, 'message': 'Failed to save DOCX document'}), 500
            else:
                with open(file_path, 'w', encoding='utf-8', errors='ignore') as f:
                    f.write(text)
        elif kind == 'table':
            columns = data.get('columns') or []
            rows = data.get('rows') or []
            table_format = data.get('format') or 'csv'
            df = pd.DataFrame(rows, columns=columns)
            if table_format == 'csv':
                df.to_csv(file_path, index=False, encoding='utf-8')
            elif table_format == 'xlsx':
                # Use pandas Excel writer
                with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                    df.to_excel(writer, index=False)
            else:
                return jsonify({'success': False, 'message': 'Unsupported table format'}), 400
        else:
            return jsonify({'success': False, 'message': 'Unsupported edit kind'}), 400

        try:
            log_user_action(current_user.id, 'edit_document', document.id, f"Edited file {document.filename}")
        except Exception:
            pass
        return jsonify({'success': True, 'message': 'Document saved successfully'})
    except Exception as e:
        app.logger.error(f"Failed to save edits: {e}")
        return jsonify({'success': False, 'message': 'Failed to save document'}), 500


@app.route('/document/<int:doc_id>/pdf')
@login_required
def document_pdf(doc_id):
    """Serve the document as a PDF inline. If the document is already a PDF, stream it.
    If it's a Word doc and LibreOffice is available, convert to PDF on-the-fly.
    """
    document = Document.query.get_or_404(doc_id)

    # Permission check
    if current_user.has_role('Student') and document.owner_id != current_user.id:
        abort(403)

    # Admins can always open; others require temporary 'open' grant
    if not current_user.can_access_admin():
        if not _has_temp_grant(current_user.id, document.id, 'open'):
            flash('Opening requires administrator approval. A request is needed.', 'warning')
            return redirect(url_for('document_detail', doc_id=document.id))

    # Resolve file path similar to view_document
    file_path = document.file_path or ''
    if not os.path.exists(file_path):
        alt = os.path.join(app.config.get('UPLOAD_FOLDER', 'uploads'), file_path)
        if os.path.exists(alt):
            file_path = alt
        else:
            candidate = os.path.join(app.config.get('UPLOAD_FOLDER', 'uploads'), document.filename or '')
            if os.path.exists(candidate):
                file_path = candidate
            else:
                app.logger.error(f"File not found for document {document.id}: {document.file_path}")
                abort(404)

    mimetype = (document.file_type or mimetypes.guess_type(file_path)[0] or '').lower()

    # If already PDF, serve inline
    if 'pdf' in mimetype or file_path.lower().endswith('.pdf'):
        resp = make_response(send_file(file_path, as_attachment=False, mimetype='application/pdf'))
        resp.headers['Content-Disposition'] = f'inline; filename="{document.filename or "document.pdf"}"'
        return resp

    # If Word doc, attempt conversion using LibreOffice helper
    if file_path.lower().endswith(('.doc', '.docx', '.odt')):
        pdf_path = _convert_word_to_pdf(file_path)
        if not pdf_path:
            # Show a minimal inline message inside the iframe instead of redirecting the whole app
            return (
                """
<html>
<head><title>Preview not available</title>
<meta name=viewport content="width=device-width, initial-scale=1" />
<style>body{font-family:system-ui, -apple-system, Segoe UI, Roboto, Arial, sans-serif;background:#0b0f19;color:#f8f9fa;margin:0;padding:16px} .card{background:#192132;border-radius:8px;padding:16px} a{color:#6ea8fe}</style>
</head>
<body>
  <div class="card">
    <h3 style="margin:0 0 8px 0;">Preview not available</h3>
    <p style="margin:0 0 6px 0;">The server couldn't convert this Word document to PDF right now.</p>
    <p style="margin:0 0 6px 0;">You can try downloading the file instead.</p>
    <p style="margin:0"><a href="%s" target="_blank" rel="noopener">Download the file</a></p>
  </div>
</body>
</html>
""" % url_for('download_document', doc_id=document.id), 200, {'Content-Type': 'text/html'}
            )

        @after_this_request
        def _cleanup_temp(response):
            try:
                if pdf_path and os.path.exists(pdf_path):
                    os.remove(pdf_path)
            except Exception:
                pass
            return response

        base = os.path.splitext(os.path.basename(file_path))[0]
        resp = make_response(send_file(pdf_path, as_attachment=False, mimetype='application/pdf'))
        resp.headers['Content-Disposition'] = f'inline; filename="{base}.pdf"'
        return resp

    # For other types (images), convert image -> single-page PDF using reportlab (if available)
    if file_path.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp')):
        try:
            from PIL import Image
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            tmp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            img = Image.open(file_path)
            w, h = img.size
            # Create PDF with the image sized to page
            c = canvas.Canvas(tmp_pdf.name, pagesize=letter)
            # Fit image to width
            max_w, max_h = letter
            ratio = min(max_w / w, max_h / h)
            new_w, new_h = w * ratio, h * ratio
            c.drawImage(file_path, 0, max_h - new_h, width=new_w, height=new_h)
            c.showPage()
            c.save()
            resp = make_response(send_file(tmp_pdf.name, as_attachment=False, mimetype='application/pdf'))
            resp.headers['Content-Disposition'] = f'inline; filename="{os.path.splitext(document.filename)[0]}.pdf"'
            return resp
        except Exception as e:
            app.logger.error(f'Image to PDF conversion failed: {e}')
            return redirect(url_for('download_document', doc_id=document.id))


@app.route('/public/view/<string:token>')
def public_document_view(token: str):
    """Serve a document inline using a short-lived signed token. No auth required.
    Token encodes (doc_id, expires). Only streams the original file (no edits/prints/download logs).
    Intended for cloud viewers (Office/Google) and 'Open in Word' deep links.
    """
    doc_id = _verify_public_token(token)
    if not doc_id:
        abort(403)
    document = Document.query.get_or_404(doc_id)

    # Resolve file path
    file_path = document.file_path or ''
    if not os.path.exists(file_path):
        alt = os.path.join(app.config.get('UPLOAD_FOLDER', 'uploads'), file_path)
        if os.path.exists(alt):
            file_path = alt
        else:
            candidate = os.path.join(app.config.get('UPLOAD_FOLDER', 'uploads'), document.filename or '')
            if os.path.exists(candidate):
                file_path = candidate
            else:
                abort(404)

    mimetype = (document.file_type or mimetypes.guess_type(file_path)[0] or 'application/octet-stream')
    resp = make_response(send_file(file_path, as_attachment=False, mimetype=mimetype))
    # Force inline with a safe filename
    disp_name = document.filename or os.path.basename(file_path)
    resp.headers['Content-Disposition'] = f'inline; filename="{disp_name}"'
    # Short cache to help external viewers load the asset during token window
    resp.headers['Cache-Control'] = 'public, max-age=300'
    resp.headers['X-Content-Type-Options'] = 'nosniff'
    return resp

    # Fallback: send the original file for download
    return redirect(url_for('download_document', doc_id=document.id))
@app.route('/document/<int:doc_id>/delete', methods=['POST'])
@login_required
def delete_document(doc_id):
    document = Document.query.get_or_404(doc_id)

    # Admin-only deletion
    if not current_user.can_access_admin():
        abort(403)

    # Remove file from disk
    try:
        if document.file_path and os.path.exists(document.file_path):
            os.remove(document.file_path)
    except Exception as e:
        app.logger.warning(f"Failed removing file for document {document.id}: {e}")

    # Remove QR image from static folder
    try:
        if document.qr_code_path:
            qr_abs = document.qr_code_path if os.path.isabs(document.qr_code_path) else os.path.join(current_app.static_folder, document.qr_code_path)
            if os.path.exists(qr_abs):
                os.remove(qr_abs)
    except Exception as e:
        app.logger.warning(f"Failed removing QR for document {document.id}: {e}")

    # Delete related access logs first to satisfy FK constraints
    try:
        AccessLog.query.filter_by(document_id=document.id).delete(synchronize_session=False)
    except Exception as e:
        app.logger.warning(f"Failed purging access logs for document {document.id}: {e}")

    title = document.title
    try:
        db.session.delete(document)
        db.session.commit()
        try:
            log_user_action(current_user.id, 'delete_document', None, f"Deleted document {title} ({doc_id})")
        except Exception:
            pass
        flash('Document deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error deleting document {doc_id}: {e}")
        flash('Failed to delete document. See logs for details.', 'error')

    return redirect(url_for('documents'))


@app.route('/document/<int:doc_id>/update', methods=['GET', 'POST'])
@login_required
def update_document(doc_id):
    document = Document.query.get_or_404(doc_id)

    # Permission: admin or original uploader with temporary update grant
    if current_user.can_access_admin():
        pass
    elif document.uploaded_by_id == current_user.id:
        if not _has_temp_grant(current_user.id, document.id, 'update'):
            flash('Updating this file requires administrator approval. Please request update access.', 'warning')
            return redirect(url_for('document_detail', doc_id=document.id))
    else:
        abort(403)

    if request.method == 'POST':
        file = request.files.get('file')
        if not file or not allowed_file(file.filename):
            flash('Please choose a valid file (pdf, doc, docx, txt, jpg, jpeg, png, gif).', 'error')
            return redirect(request.url)

        from werkzeug.utils import secure_filename as _secure
        filename = _secure(file.filename)
        unique_filename = f"{uuid.uuid4()}_{filename}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        try:
            file.save(file_path)
        except Exception as e:
            app.logger.error(f'Failed saving uploaded file: {e}')
            flash('Failed to save the uploaded file.', 'error')
            return redirect(request.url)

        # Optionally remove the old file from disk to free space
        try:
            if document.file_path and os.path.exists(document.file_path):
                os.remove(document.file_path)
        except Exception:
            pass

        # Update document file metadata
        document.filename = filename
        document.file_path = file_path
        try:
            document.file_size = os.path.getsize(file_path)
        except Exception:
            document.file_size = None
        document.file_type = file.content_type
        try:
            db.session.commit()
            log_user_action(current_user.id, 'update_file', document.id, f"Updated file to {filename}")
            flash('Document file updated successfully.', 'success')
        except Exception as e:
            db.session.rollback()
            app.logger.error(f'Document update failed: {e}')
            flash('Failed to update document.', 'error')
            return redirect(request.url)

        return redirect(url_for('document_detail', doc_id=document.id))

    return render_template('update_document.html', document=document)

@app.route('/document/<int:doc_id>/owner_details/update', methods=['POST'])
@login_required
def update_owner_details(doc_id):
    document = Document.query.get_or_404(doc_id)

    # Permission: Admins can always edit; others require temporary 'owner_edit' grant
    if not current_user.can_access_admin():
        if not _has_temp_grant(current_user.id, document.id, 'owner_edit'):
            flash('Editing owner information requires administrator approval. Please request access.', 'warning')
            return redirect(url_for('document_detail', doc_id=doc_id))

    fn = (request.form.get('owner_first_name') or '').strip()
    mn = (request.form.get('owner_middle_name') or '').strip()
    ln = (request.form.get('owner_last_name') or '').strip()
    addr = (request.form.get('owner_address') or '').strip()

    try:
        od = document.owner_details
    except Exception:
        od = None

    try:
        if not od:
            od = OwnerDocumentDetails(document_id=document.id)
            db.session.add(od)
        od.first_name = fn or None
        od.middle_name = mn or None
        od.last_name = ln or None
        od.address = addr or None
        db.session.commit()
        try:
            log_user_action(current_user.id, 'update_owner_details', document.id, 'Updated owner information')
        except Exception:
            pass
        flash('Owner information updated successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        app.logger.error(f'update_owner_details failed for doc {doc_id}: {e}')
        flash('Failed to update owner information.', 'error')

    return redirect(url_for('document_detail', doc_id=doc_id))

@app.route('/document/<int:doc_id>/request/<string:kind>', methods=['POST'])
@login_required
def request_access(doc_id, kind):
    document = Document.query.get_or_404(doc_id)
    kind = (kind or '').strip().lower()
    valid_kinds = {'download', 'print', 'update', 'open', 'email', 'owner_edit'}
    if kind not in valid_kinds:
        abort(400)

    # Basic permission: students can't request actions on others' docs
    if current_user.has_role('Student') and document.owner_id != current_user.id:
        abort(403)

    # Prevent duplicate pending requests by the same user for the same kind
    try:
        req_action = f'request_{kind}'
        grant_action = f'{kind}_grant'
        reject_action = f'{req_action}_reject'
        last_req = AccessLog.query.filter_by(user_id=current_user.id, document_id=document.id, action=req_action).order_by(desc(AccessLog.timestamp)).first()
        if last_req:
            resolution = AccessLog.query.filter(
                AccessLog.user_id == current_user.id,
                AccessLog.document_id == document.id,
                AccessLog.action.in_([grant_action, reject_action]),
                AccessLog.timestamp >= last_req.timestamp
            ).first()
            if resolution is None:
                flash('A request for this action is already pending approval.', 'info')
                return redirect(url_for('document_detail', doc_id=doc_id))
    except Exception:
        pass

    # Log the request
    try:
        details = f"{kind.capitalize()} access requested."
        log_user_action(current_user.id, f'request_{kind}', document.id, details)
    except Exception:
        pass

    # Email admins
    try:
        admins = User.query.join(Role).filter(Role.name == 'Admin').all()
        recipients = [u.email for u in admins if u.email]
        if recipients:
            subject = f"{kind.capitalize()} access request: {document.document_number or ('#' + str(document.id))} - {document.title}"
            view_link = url_for('document_detail', doc_id=document.id, _external=True)
            approve_hint = 'Approve in Admin panel activity logs.'
            body = (
                f"User: {current_user.get_full_name()} (@{current_user.username})\n"
                f"Document: {document.title}\n"
                f"Document ID: {document.document_number or ('#' + str(document.id))}\n"
                f"Request: {kind}\n\n"
                f"View: {view_link}\n\n{approve_hint}\n"
            )
            send_email(subject, recipients, body)
    except Exception as e:
        app.logger.warning(f"Failed to send {kind} request email: {e}")

    flash('Your request has been sent to the administrators.', 'success')
    return redirect(url_for('document_detail', doc_id=doc_id))

@app.route('/admin/access_requests/<int:request_id>/approve', methods=['POST'])
@login_required
def approve_access_request(request_id):
    if not current_user.can_access_admin():
        return jsonify({'success': False, 'message': 'Access denied'}), 403

    log = AccessLog.query.get_or_404(request_id)
    if not log.action or not log.action.startswith('request_'):
        return jsonify({'success': False, 'message': 'Invalid request'}), 400
    kind = log.action.replace('request_', '', 1)
    grant_action = f"{kind}_grant"
    try:
        db.session.add(AccessLog(
            user_id=log.user_id,
            document_id=log.document_id,
            action=grant_action,
            ip_address=request.remote_addr,
            user_agent=request.headers.get('User-Agent'),
            details=f"Temporary {kind} access granted for 30 minutes"
        ))
        db.session.commit()
        return jsonify({'success': True, 'message': f'{kind.capitalize()} access granted for 30 minutes.'})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"approve_access_request failed: {e}")
        return jsonify({'success': False, 'message': 'Failed to approve request'}), 500

@app.route('/admin/access_requests/<int:request_id>/reject', methods=['POST'])
@login_required
def reject_access_request(request_id):
    if not current_user.can_access_admin():
        return jsonify({'success': False, 'message': 'Access denied'}), 403

    log = AccessLog.query.get_or_404(request_id)
    if not log.action or not log.action.startswith('request_'):
        return jsonify({'success': False, 'message': 'Invalid request'}), 400
    try:
        db.session.add(AccessLog(
            user_id=log.user_id,
            document_id=log.document_id,
            action=f"{log.action}_reject",
            ip_address=request.remote_addr,
            user_agent=request.headers.get('User-Agent'),
            details=f"Rejected {log.action} #{request_id}"
        ))
        db.session.commit()
        return jsonify({'success': True, 'message': 'Request rejected.'})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"reject_access_request failed: {e}")
        return jsonify({'success': False, 'message': 'Failed to reject request'}), 500

@app.route('/admin/edit_requests/reject_all', methods=['POST'])
@login_required
def reject_all_edit_requests():
    if not current_user.can_access_admin():
        return jsonify({'success': False, 'message': 'Access denied'}), 403
    try:
        # Find unresolved edit requests (no later grant or reject recorded)
        raw_requests = AccessLog.query.filter_by(action='request_edit').order_by(AccessLog.timestamp.asc()).all()
        rejected = 0
        for req in raw_requests:
            has_resolution = AccessLog.query.filter(
                AccessLog.document_id == req.document_id,
                AccessLog.user_id == req.user_id,
                AccessLog.action.in_(['edit_grant', 'request_edit_reject']),
                AccessLog.timestamp >= req.timestamp
            ).first()
            if has_resolution:
                continue
            db.session.add(AccessLog(
                user_id=req.user_id,
                document_id=req.document_id,
                action='request_edit_reject',
                ip_address=request.remote_addr,
                user_agent=request.headers.get('User-Agent'),
                details='Bulk reject (admin): cleared pending edit request'
            ))
            rejected += 1
        db.session.commit()
        return jsonify({'success': True, 'rejected': rejected})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"reject_all_edit_requests failed: {e}")
        return jsonify({'success': False, 'message': 'Failed to bulk reject edit requests'}), 500

@app.route('/admin/access_requests/reject_all', methods=['POST'])
@login_required
def reject_all_access_requests():
    if not current_user.can_access_admin():
        return jsonify({'success': False, 'message': 'Access denied'}), 403
    try:
        valid_actions = ['request_download', 'request_print', 'request_update', 'request_open', 'request_email', 'request_owner_edit']
        raw_access = AccessLog.query.filter(AccessLog.action.in_(valid_actions)).order_by(AccessLog.timestamp.asc()).all()
        rejected = 0
        for req in raw_access:
            kind = (req.action or '').replace('request_', '')
            resolution = AccessLog.query.filter(
                AccessLog.document_id == req.document_id,
                AccessLog.user_id == req.user_id,
                AccessLog.action.in_([f'{kind}_grant', f'request_{kind}_reject']),
                AccessLog.timestamp >= req.timestamp
            ).first()
            if resolution:
                continue
            db.session.add(AccessLog(
                user_id=req.user_id,
                document_id=req.document_id,
                action=f'request_{kind}_reject',
                ip_address=request.remote_addr,
                user_agent=request.headers.get('User-Agent'),
                details='Bulk reject (admin): cleared pending access request'
            ))
            rejected += 1
        db.session.commit()
        return jsonify({'success': True, 'rejected': rejected})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"reject_all_access_requests failed: {e}")
        return jsonify({'success': False, 'message': 'Failed to bulk reject access requests'}), 500

@app.route('/document/<int:doc_id>/request_edit', methods=['POST'])
@login_required
def request_edit_access(doc_id):
    document = Document.query.get_or_404(doc_id)

    # If admin, no need to request
    if current_user.can_access_admin():
        flash('You already have edit access as an admin.', 'info')
        return redirect(url_for('edit_document', doc_id=doc_id))

    reason = (request.form.get('reason') or '').strip()

    # Log the request in audit trail
    try:
        details = f"Edit access requested. Reason: {reason[:500]}"
        log_user_action(current_user.id, 'request_edit', document.id, details)
    except Exception:
        pass

    # Notify admins via email if possible
    try:
        admins = User.query.join(Role).filter(Role.name == 'Admin').all()
        recipients = [u.email for u in admins if u.email]
        if recipients:
            subject = f"Edit access request: {document.document_number or ('#' + str(document.id))} - {document.title}"
            view_link = url_for('document_detail', doc_id=document.id, _external=True)
            edit_link = url_for('edit_document', doc_id=document.id, _external=True)
            body = (
                f"User: {current_user.get_full_name()} (@{current_user.username})\n"
                f"Document: {document.title}\n"
                f"Document ID: {document.document_number or ('#' + str(document.id))}\n"
                f"Reason: {reason}\n\n"
                f"View: {view_link}\n"
                f"Admin Edit: {edit_link}\n"
            )
            send_email(subject, recipients, body)
    except Exception as e:
        app.logger.warning(f"Failed to send edit request email: {e}")

    flash('Your edit request has been sent to the administrators.', 'success')
    return redirect(url_for('document_detail', doc_id=doc_id))


@app.route('/admin/edit_requests/<int:request_id>/approve', methods=['POST'])
@login_required
def approve_edit_request(request_id):
    if not current_user.can_access_admin():
        return jsonify({'success': False, 'message': 'Access denied'}), 403

    log = AccessLog.query.get_or_404(request_id)
    if log.action != 'request_edit':
        return jsonify({'success': False, 'message': 'Invalid request'}), 400

    try:
        grant_details = 'Temporary edit access granted for 30 minutes'
        db.session.add(AccessLog(
            user_id=log.user_id,
            document_id=log.document_id,
            action='edit_grant',
            ip_address=request.remote_addr,
            user_agent=request.headers.get('User-Agent'),
            details=grant_details
        ))
        db.session.commit()
        requester = User.query.get(log.user_id)
        document = Document.query.get(log.document_id)
        if requester and requester.email and document:
            try:
                edit_link = url_for('edit_document', doc_id=document.id, _external=True)
                subject = 'Edit access approved'
                body = (
                    f"Your request to edit '{document.title}' has been approved for 30 minutes.\n\n"
                    f"Open editor: {edit_link}\n"
                )
                send_email(subject, [requester.email], body)
            except Exception:
                pass
        return jsonify({'success': True, 'message': 'Edit access granted for 24 hours.'})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"approve_edit_request failed: {e}")
        return jsonify({'success': False, 'message': 'Failed to approve request'}), 500


@app.route('/admin/edit_requests/<int:request_id>/reject', methods=['POST'])
@login_required
def reject_edit_request(request_id):
    if not current_user.can_access_admin():
        return jsonify({'success': False, 'message': 'Access denied'}), 403

    log = AccessLog.query.get_or_404(request_id)
    if log.action != 'request_edit':
        return jsonify({'success': False, 'message': 'Invalid request'}), 400

    try:
        db.session.add(AccessLog(
            user_id=log.user_id,
            document_id=log.document_id,
            action='request_edit_reject',
            ip_address=request.remote_addr,
            user_agent=request.headers.get('User-Agent'),
            details=f"Rejected request #{request_id}"
        ))
        db.session.commit()
        requester = User.query.get(log.user_id)
        document = Document.query.get(log.document_id)
        if requester and requester.email and document:
            try:
                subject = 'Edit access request rejected'
                body = (
                    f"Your request to edit '{document.title}' was rejected by an administrator.\n"
                )
                send_email(subject, [requester.email], body)
            except Exception:
                pass
        return jsonify({'success': True, 'message': 'Request rejected.'})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"reject_edit_request failed: {e}")
        return jsonify({'success': False, 'message': 'Failed to reject request'}), 500


@app.route('/document/<int:doc_id>/email', methods=['POST'])
@login_required
def send_document_email(doc_id):
    document = Document.query.get_or_404(doc_id)

    # Permission: Admins can always send. Others require temporary 'email' grant.
    if not current_user.can_access_admin():
        if not _has_temp_grant(current_user.id, document.id, 'email'):
            flash('Emailing this document requires administrator approval. Please request email access.', 'warning')
            return redirect(url_for('document_detail', doc_id=doc_id))

    recipients_raw = request.form.get('recipients')
    subject = request.form.get('subject') or f"Document shared: {document.title}"
    message = request.form.get('message', '')
    attach = request.form.get('attach', 'no') == 'yes'
    include_details = request.form.get('include_details', 'no') == 'yes'

    if not recipients_raw:
        flash('At least one recipient email is required.', 'error')
        return redirect(url_for('document_detail', doc_id=doc_id))

    # Parse comma/semicolon-separated recipient list
    recipients = [r.strip() for r in re.split(r'[,;\s]+', recipients_raw) if r.strip()]

    # Basic email validation
    valid_recipients = []
    invalid = []
    for r in recipients:
        if re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", r):
            valid_recipients.append(r)
        else:
            invalid.append(r)

    if not valid_recipients:
        flash('No valid recipient emails provided.', 'error')
        return redirect(url_for('document_detail', doc_id=doc_id))

    # Prefer a public landing URL so recipients can open the document without logging in
    try:
        if getattr(document, 'document_number', None):
            link = url_for('qr_landing_by_code', code=document.document_number, _external=True)
        else:
            link = url_for('qr_landing', doc_id=document.id, _external=True)
    except Exception:
        link = request.url_root.rstrip('/') + url_for('document_detail', doc_id=document.id)
    
    # Compose a formal letter body
    sender_name = current_user.get_full_name()
    sender_email = current_user.email
    date = datetime.now().strftime('%B %d, %Y')
    recipient_name = "Recipient"  # Placeholder since recipient details are not available
    document_list = document.title
    reason = message or "application requirements"
    details = ""
    if include_details:
        details = f"\n\nDocument Details:\nDocument ID: {document.document_number or ('#' + str(document.id))}\nTitle: {document.title}\nOwner: {document.owner.get_full_name()}\nUploaded by: {document.uploader.get_full_name()}\n"

    body = f"""{sender_name}
{sender_email}
{date}

{recipient_name}

Dear {recipient_name},

I am submitting the following documents for your review and reference:

{document_list}

These documents are being sent in compliance with {reason}. Kindly acknowledge receipt at your earliest convenience.

You can view the document here: {link}{details}

Should you need any additional information, please do not hesitate to contact me.

Thank you for your attention.

Sincerely,
{sender_name}
"""

    attachments = None
    if attach:
        if document.file_path and os.path.exists(document.file_path):
            mimetype = document.file_type or mimetypes.guess_type(document.file_path)[0] or 'application/octet-stream'
            attachments = [(document.file_path, document.filename, mimetype)]

    success = send_email(subject, valid_recipients, body, attachments=attachments, reply_to=current_user.email)

    if success:
        log_user_action(current_user.id, 'email', document.id, f"Sent document to {', '.join(valid_recipients)}")
        msg = 'Email sent successfully.'
        if invalid:
            msg += f" Note: some addresses were invalid and were skipped: {', '.join(invalid)}"
        flash(msg, 'success')
    else:
        # Provide a helpful hint: SMTP may not be configured. Check both MAIL_* and EMAIL_* keys.
        cfg = current_app.config
        smtp_host = cfg.get('MAIL_SERVER') or cfg.get('EMAIL_HOST')
        smtp_port = cfg.get('MAIL_PORT') or cfg.get('EMAIL_PORT')
        smtp_user = cfg.get('MAIL_USERNAME') or cfg.get('EMAIL_HOST_USER')
        smtp_pass_present = bool(cfg.get('MAIL_PASSWORD') or cfg.get('EMAIL_HOST_PASSWORD'))
        has_credentials = bool(smtp_user and smtp_pass_present)
        current_app.logger.warning(f"send_document_email: SMTP send failed. SMTP host: {smtp_host}, port: {smtp_port}, has_credentials: {has_credentials}")
        flash('Failed to send email. Check server configuration (MAIL_SERVER/MAIL_PORT/MAIL_USERNAME or EMAIL_HOST/EMAIL_PORT/EMAIL_HOST_USER) and application logs. You can also use the "Open in Email App" button in the Email dialog to send this message from your own mail client.', 'error')

    return redirect(url_for('document_detail', doc_id=doc_id))

# @app.route('/document/<int:doc_id>/email', methods=['POST'])
# @login_required
# def send_document_email(doc_id):
#     document = Document.query.get_or_404(doc_id)

#     # Permission: only owner, uploader, or admin roles can send
#     if current_user.has_role('Student') and document.owner_id != current_user.id and document.uploaded_by_id != current_user.id:
#         return jsonify({'error': 'Access denied'}), 403

#     recipients_raw = request.form.get('recipients')
#     subject = request.form.get('subject') or f"Document shared: {document.title}"
#     message = request.form.get('message', '')
#     attach = request.form.get('attach', 'no') == 'yes'
#     include_details = request.form.get('include_details', 'no') == 'yes'

#     if not recipients_raw:
#         flash('At least one recipient email is required.', 'error')
#         return redirect(url_for('document_detail', doc_id=doc_id))

#     # Parse comma/semicolon-separated recipient list
#     recipients = [r.strip() for r in re.split(r'[,;\s]+', recipients_raw) if r.strip()]

#     # Basic email validation
#     valid_recipients = []
#     invalid = []
#     for r in recipients:
#         if re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", r):
#             valid_recipients.append(r)
#         else:
#             invalid.append(r)

#     if not valid_recipients:
#         flash('No valid recipient emails provided.', 'error')
#         return redirect(url_for('document_detail', doc_id=doc_id))

#     # Prefer a public landing URL so recipients can open the document without logging in
#     try:
#         if getattr(document, 'document_number', None):
#             link = url_for('qr_landing_by_code', code=document.document_number, _external=True)
#         else:
#             link = url_for('qr_landing', doc_id=document.id, _external=True)
#     except Exception:
#         link = request.url_root.rstrip('/') + url_for('document_detail', doc_id=document.id)
#     # Compose a helpful body including who/what when requested
#     details = f"Document ID: {document.document_number or ('#' + str(document.id))}\nTitle: {document.title}\nOwner: {document.owner.get_full_name()}\nUploaded by: {document.uploader.get_full_name()}\n"
#     body_lines = []
#     if message:
#         body_lines.append(message)
#     if include_details:
#         body_lines.append('\n'.join(['', 'Document Details:', details]))
#     body_lines.append(f"You can view the document here: {link}")
#     body_lines.append(f"\nSent by {current_user.get_full_name()}")
#     body = "\n\n".join(body_lines)

#     attachments = None
#     if attach:
#         if document.file_path and os.path.exists(document.file_path):
#             mimetype = document.file_type or mimetypes.guess_type(document.file_path)[0] or 'application/octet-stream'
#             attachments = [(document.file_path, document.filename, mimetype)]

#     success = send_email(subject, valid_recipients, body, attachments=attachments, reply_to=current_user.email)

#     if success:
#         log_user_action(current_user.id, 'email', document.id, f"Sent document to {', '.join(valid_recipients)}")
#         msg = 'Email sent successfully.'
#         if invalid:
#             msg += f" Note: some addresses were invalid and were skipped: {', '.join(invalid)}"
#         flash(msg, 'success')
#     else:
#         # Provide a helpful hint: SMTP may not be configured. Log minimal config presence for debugging (no secrets)
#         cfg = current_app.config
#         smtp_host = cfg.get('EMAIL_HOST', None)
#         smtp_port = cfg.get('EMAIL_PORT', None)
#         has_credentials = bool(cfg.get('EMAIL_HOST_USER') and cfg.get('EMAIL_HOST_PASSWORD'))
#         current_app.logger.warning(f"send_document_email: SMTP send failed. SMTP host: {smtp_host}, port: {smtp_port}, has_credentials: {has_credentials}")
#         flash('Failed to send email. Check server configuration (EMAIL_HOST/EMAIL_PORT/EMAIL_HOST_USER) and application logs. You can also use the "Open in Email App" button in the Email dialog to send this message from your own mail client.', 'error')

#     return redirect(url_for('document_detail', doc_id=doc_id))


@app.route('/document/<int:doc_id>/regenerate_qr', methods=['POST'])
@login_required
def regenerate_qr(doc_id):
    """Regenerate the QR image for a single document (admin or uploader only)."""
    document = Document.query.get_or_404(doc_id)

    if not (current_user.can_access_admin() or document.uploaded_by_id == current_user.id):
        abort(403)

    # Delete old QR file if present
    try:
        if document.qr_code_path:
            old_path = os.path.join(current_app.static_folder, document.qr_code_path)
            if os.path.exists(old_path):
                os.remove(old_path)
    except Exception as e:
        app.logger.warning(f"Failed removing old QR: {e}")

    # Generate new QR (this will annotate with id and details if PIL present)
    try:
        new_rel = generate_qr_code(document.id)
        document.qr_code_path = new_rel
        db.session.commit()
        log_user_action(current_user.id, 'regenerate_qr', document.id, 'Regenerated QR code')
        flash('QR code regenerated.', 'success')
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error regenerating QR for {document.id}: {e}")
        flash('Failed to regenerate QR code. See logs.', 'error')

    return redirect(url_for('document_detail', doc_id=doc_id))


@app.route('/document/<int:doc_id>/qr/download', methods=['GET'])
@login_required
def download_qr(doc_id):
    """Download the QR image for a document as an attachment, regenerating if missing."""
    document = Document.query.get_or_404(doc_id)

    # Permission: students can only download QR for their own documents
    if current_user.has_role('Student') and document.owner_id != current_user.id:
        abort(403)

    # Resolve QR path (relative to static) and regenerate if necessary
    qr_rel = document.qr_code_path
    qr_abs = None
    if qr_rel:
        qr_abs = qr_rel if os.path.isabs(qr_rel) else os.path.join(current_app.static_folder, qr_rel)

    if not qr_rel or not os.path.exists(qr_abs):
        try:
            new_rel = generate_qr_code(document.id)
            document.qr_code_path = new_rel
            db.session.commit()
            qr_rel = new_rel
            qr_abs = new_rel if os.path.isabs(new_rel) else os.path.join(current_app.static_folder, new_rel)
        except Exception as e:
            db.session.rollback()
            app.logger.error(f"Failed to generate QR for download: {e}")

    if not qr_abs or not os.path.exists(qr_abs):
        abort(404)

    # Build a friendly download name
    try:
        code = document.document_number or f"DOC{document.id:06d}"
    except Exception:
        code = f"DOC{document.id}"
    download_name = f"qr_{code}.png"

    # Log and send file
    try:
        log_user_action(current_user.id, 'download_qr', document.id, f"Downloaded QR for {code}")
    except Exception:
        pass

    return send_file(qr_abs, as_attachment=True, download_name=download_name, mimetype='image/png')


@app.route('/admin/regenerate_qr_all', methods=['POST'])
@login_required
def regenerate_qr_all():
    """Admin-only: regenerate QR images for all documents. Use with caution."""
    if not current_user.can_access_admin():
        abort(403)

    count = 0
    try:
        docs = Document.query.all()
        for doc in docs:
            try:
                # remove old
                if doc.qr_code_path:
                    old = os.path.join(current_app.static_folder, doc.qr_code_path)
                    if os.path.exists(old):
                        os.remove(old)
                new_rel = generate_qr_code(doc.id)
                doc.qr_code_path = new_rel
                count += 1
            except Exception as e:
                app.logger.warning(f"Skipping QR regen for {doc.id}: {e}")
        db.session.commit()
        flash(f'Regenerated QR for {count} documents.', 'success')
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Bulk QR regeneration failed: {e}")
        flash('Bulk regeneration failed. See logs.', 'error')

    return redirect(url_for('admin_panel'))

@app.route('/scan_qr')
@login_required
def scan_qr():
    return render_template('scan_qr.html')

@app.route('/qr_lookup/<int:doc_id>')
@login_required
def qr_lookup(doc_id):
    document = Document.query.get_or_404(doc_id)
    
    # Check permissions
    if current_user.has_role('Student') and document.owner_id != current_user.id:
        return jsonify({'error': 'Access denied'}), 403
    
    # Log the QR access
    log_user_action(current_user.id, 'qr_scan', doc_id, f"Accessed document via QR: {document.title}")
    document.increment_access()
    
    return jsonify({
        'id': document.id,
        'document_number': document.document_number,
        'title': document.title,
        'type': document.document_type.name,
        'owner': document.owner.get_full_name(),
        'created_at': document.created_at.strftime('%Y-%m-%d'),
        'download_url': url_for('download_document', doc_id=document.id),
        'view_url': url_for('document_detail', doc_id=document.id)
    })


@app.route('/qr_lookup/code/<string:code>')
@login_required
def qr_lookup_by_code(code):
    # Find document by its human-friendly code (document_number) or numeric fallback
    document = Document.query.filter_by(document_number=code).first()
    if not document:
        # try numeric id fallback
        try:
            did = int(code)
            document = Document.query.get(did)
        except Exception:
            document = None

    if not document:
        return jsonify({'error': 'Not found'}), 404

    # Permission check
    if current_user.has_role('Student') and document.owner_id != current_user.id:
        return jsonify({'error': 'Access denied'}), 403

    # Log the QR access
    log_user_action(current_user.id, 'qr_scan', document.id, f"Accessed document via QR code: {document.title}")
    document.increment_access()

    return jsonify({
        'id': document.id,
        'document_number': document.document_number,
        'title': document.title,
        'type': document.document_type.name,
        'owner': document.owner.get_full_name(),
        'created_at': document.created_at.strftime('%Y-%m-%d'),
        'download_url': url_for('download_document', doc_id=document.id),
        'view_url': url_for('document_detail', doc_id=document.id)
    })


# Public landing page for QR scans from mobile cameras. Shows minimal info (ID + title).
@app.route('/qr/<int:doc_id>')
def qr_landing(doc_id):
    document = Document.query.get_or_404(doc_id)
    # Show only minimal info for public QR landing
    return render_template('qr_landing.html', document=document)


@app.route('/qr/code/<string:code>')
def qr_landing_by_code(code):
    # find by document_number or numeric id fallback
    document = Document.query.filter_by(document_number=code).first()
    if not document:
        try:
            did = int(code)
            document = Document.query.get(did)
        except Exception:
            document = None
    if not document:
        return render_template('qr_landing.html', document=None, not_found=True), 404
    return render_template('qr_landing.html', document=document)

@app.route('/analytics')
@login_required
def analytics():
    if not (current_user.has_role('Admin') or current_user.has_role('Teacher') or current_user.has_role('Registrar')):
        flash('You do not have permission to access analytics.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get analytics data
    stats = get_document_stats()
    predictions = predict_document_requests()
    
    return render_template('analytics.html', stats=stats, predictions=predictions)

@app.route('/api/analytics/chart_data')
@login_required
def analytics_chart_data():
    if not (current_user.has_role('Admin') or current_user.has_role('Teacher') or current_user.has_role('Registrar')):
        return jsonify({'error': 'Access denied'}), 403
    
    chart_type = request.args.get('type', 'monthly_requests')
    
    if chart_type == 'monthly_requests':
        # Get monthly document requests for the last 12 months
        start_date = datetime.utcnow() - timedelta(days=365)
        data = db.session.query(
            func.date_trunc('month', AccessLog.timestamp).label('month'),
            func.count(AccessLog.id).label('count')
        ).filter(
            AccessLog.timestamp >= start_date,
            AccessLog.action.in_(['view', 'download'])
        ).group_by('month').order_by('month').all()
        
        return jsonify({
            'labels': [row.month.strftime('%Y-%m') for row in data],
            'data': [row.count for row in data]
        })
    
    elif chart_type == 'document_types':
        # Get document type distribution
        data = db.session.query(
            DocumentType.name,
            func.count(Document.id).label('count')
        ).join(Document).filter(Document.is_active == True).group_by(DocumentType.name).all()
        
        return jsonify({
            'labels': [row.name for row in data],
            'data': [row.count for row in data]
        })
    
    elif chart_type == 'top_documents':
        # Get top accessed documents
        data = Document.query.filter_by(is_active=True).order_by(desc(Document.access_count)).limit(10).all()
        
        return jsonify({
            'labels': [doc.title[:30] + '...' if len(doc.title) > 30 else doc.title for doc in data],
            'data': [doc.access_count for doc in data]
        })
    
    return jsonify({'error': 'Invalid chart type'}), 400

@app.route('/admin')
@login_required
def admin_panel():
    if not current_user.can_access_admin():
        flash('You do not have permission to access the admin panel.', 'error')
        return redirect(url_for('dashboard'))
    
    # Get admin statistics
    total_users = User.query.count()
    total_documents = Document.query.filter_by(is_active=True).count()
    recent_logs = AccessLog.query.order_by(desc(AccessLog.timestamp)).limit(10).all()
    
    # Document types for admin management
    document_types = DocumentType.query.order_by(DocumentType.name).all()

    # Pending edit requests: only requests without a subsequent approval/rejection
    try:
        raw_requests = AccessLog.query.filter_by(action='request_edit').order_by(desc(AccessLog.timestamp)).limit(200).all()
        pending = []
        for req in raw_requests:
            has_resolution = AccessLog.query.filter(
                AccessLog.document_id == req.document_id,
                AccessLog.user_id == req.user_id,
                AccessLog.action.in_(['edit_grant', 'request_edit_reject']),
                AccessLog.timestamp >= req.timestamp
            ).first()
            if not has_resolution:
                pending.append(req)
        # Trim to 50 most recent unresolved
        pending_edit_requests = pending[:50]
    except Exception:
        pending_edit_requests = []

    # Pending access requests (download/print/update/open/email) without subsequent grant/reject
    try:
        raw_access = AccessLog.query.filter(AccessLog.action.in_([
            'request_download', 'request_print', 'request_update', 'request_open', 'request_email', 'request_owner_edit'
        ])).order_by(desc(AccessLog.timestamp)).limit(300).all()
        pend = []
        for req in raw_access:
            kind = req.action.replace('request_', '')
            resolution = AccessLog.query.filter(
                AccessLog.document_id == req.document_id,
                AccessLog.user_id == req.user_id,
                AccessLog.action.in_([f'{kind}_grant', f'request_{kind}_reject']),
                AccessLog.timestamp >= req.timestamp
            ).first()
            if not resolution:
                pend.append(req)
        pending_access_requests = pend[:100]
    except Exception:
        pending_access_requests = []

    # Active temporary grants (edit/download/print/update/open/email) within 30 minutes window
    try:
        cutoff = datetime.utcnow() - timedelta(minutes=30)
        active_grants = AccessLog.query.filter(
            AccessLog.action.in_(['edit_grant', 'download_grant', 'print_grant', 'update_grant', 'open_grant', 'email_grant', 'owner_edit_grant']),
            AccessLog.timestamp >= cutoff
        ).order_by(desc(AccessLog.timestamp)).limit(500).all()
        # Filter out revoked grants
        filtered = []
        for g in active_grants:
            try:
                kind = (g.action or '').replace('_grant', '')
                revoke_action = f"{kind}_revoke"
                revoked = AccessLog.query.filter(
                    AccessLog.user_id == g.user_id,
                    AccessLog.document_id == g.document_id,
                    AccessLog.action == revoke_action,
                    AccessLog.timestamp >= g.timestamp
                ).first()
                if not revoked:
                    filtered.append(g)
            except Exception:
                filtered.append(g)
        active_grants = filtered
    except Exception:
        active_grants = []

    # Build pre-formatted expiration info for template
    try:
        active_grants_info = []
        for g in active_grants:
            exp = g.timestamp + timedelta(minutes=30)
            active_grants_info.append({
                'grant': g,
                'exp_str': exp.strftime('%Y-%m-%d %H:%M:%S'),
                'exp_iso': exp.strftime('%Y-%m-%dT%H:%M:%SZ')
            })
    except Exception:
        active_grants_info = []

    return render_template('admin.html',
                         total_users=total_users,
                         total_documents=total_documents,
                         recent_logs=recent_logs,
                         document_types=document_types,
                         pending_edit_requests=pending_edit_requests,
                         pending_access_requests=pending_access_requests,
                         active_grants=active_grants,
                         active_grants_info=active_grants_info)

@app.route('/admin/grants/<int:grant_id>/revoke', methods=['POST'])
@login_required
def revoke_grant(grant_id: int):
    """Admin-only: revoke an active grant immediately by recording a {kind}_revoke entry.
    This prevents further use of the temporary grant.
    """
    if not current_user.can_access_admin():
        return jsonify({'success': False, 'message': 'Access denied'}), 403

    grant = AccessLog.query.get_or_404(grant_id)
    if not grant.action or not grant.action.endswith('_grant'):
        return jsonify({'success': False, 'message': 'Invalid grant'}), 400

    try:
        kind = grant.action.replace('_grant', '')
        revoke_action = f"{kind}_revoke"
        db.session.add(AccessLog(
            user_id=grant.user_id,  # target user of the grant
            document_id=grant.document_id,
            action=revoke_action,
            ip_address=request.remote_addr,
            user_agent=request.headers.get('User-Agent'),
            details=f"Grant revoked by admin @{current_user.username}"
        ))
        db.session.commit()
        return jsonify({'success': True, 'message': f'{kind} grant revoked.'})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"revoke_grant failed: {e}")
        return jsonify({'success': False, 'message': 'Failed to revoke grant'}), 500


@app.route('/admin/add_document_type', methods=['POST'])
@login_required
def add_document_type():
    if not current_user.can_access_admin():
        flash('You do not have permission to perform this action.', 'error')
        return redirect(url_for('admin_panel'))

    name = request.form.get('name')
    description = request.form.get('description')

    if not name or name.strip() == '':
        flash('Document type name is required.', 'error')
        return redirect(url_for('admin_panel'))

    # Prevent duplicates
    existing = DocumentType.query.filter(func.lower(DocumentType.name) == name.strip().lower()).first()
    if existing:
        flash('A document type with that name already exists.', 'error')
        return redirect(url_for('admin_panel'))

    doc_type = DocumentType(name=name.strip(), description=description)
    db.session.add(doc_type)
    db.session.commit()

    flash('Document type added successfully.', 'success')
    return redirect(url_for('admin_panel'))


@app.route('/admin/delete_document_type/<int:type_id>', methods=['POST'])
@login_required
def delete_document_type(type_id):
    if not current_user.can_access_admin():
        flash('You do not have permission to perform this action.', 'error')
        return redirect(url_for('admin_panel'))

    doc_type = DocumentType.query.get_or_404(type_id)

    # Prevent deletion if any documents use this type
    in_use = Document.query.filter_by(document_type_id=doc_type.id).first()
    if in_use:
        flash('Cannot delete document type: it is in use by existing documents.', 'error')
        return redirect(url_for('admin_panel'))

    try:
        db.session.delete(doc_type)
        db.session.commit()
        flash('Document type deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Error deleting document type {type_id}: {e}")
        flash('An error occurred while deleting the document type.', 'error')

    return redirect(url_for('admin_panel'))

@app.route('/admin/users/<int:user_id>/update', methods=['POST'])
@login_required
def admin_update_user(user_id):
    if not current_user.can_access_admin():
        return jsonify({'success': False, 'message': 'Access denied'}), 403

    user = User.query.get_or_404(user_id)

    # Parse form fields
    first_name = request.form.get('first_name', '').strip()
    last_name = request.form.get('last_name', '').strip()
    email = request.form.get('email', '').strip()
    student_id = request.form.get('student_id', '').strip() or None
    role_id = request.form.get('role_id')

    # Validation
    if not all([first_name, last_name, email, role_id]):
        return jsonify({'success': False, 'message': 'Missing required fields'}), 400

    # Unique email constraint (excluding this user)
    existing_email = User.query.filter(User.email == email, User.id != user.id).first()
    if existing_email:
        return jsonify({'success': False, 'message': 'Email already in use by another user'}), 400

    # Unique student_id when provided
    if student_id:
        existing_sid = User.query.filter(User.student_id == student_id, User.id != user.id).first()
        if existing_sid:
            return jsonify({'success': False, 'message': 'Student/Employee ID already in use by another user'}), 400

    # Ensure role exists
    role = Role.query.get(role_id)
    if not role:
        return jsonify({'success': False, 'message': 'Invalid role selected'}), 400

    # Apply updates
    try:
        user.first_name = first_name
        user.last_name = last_name
        user.email = email
        user.student_id = student_id
        user.role_id = role.id
        db.session.commit()
        try:
            log_user_action(current_user.id, 'admin_update_user', None, f"Updated user {user.username} ({user.id})")
        except Exception:
            pass
        return jsonify({'success': True, 'message': 'User updated successfully'})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"admin_update_user failed: {e}")
        return jsonify({'success': False, 'message': 'Failed to update user'}), 500


@app.route('/admin/users/<int:user_id>/reset_password', methods=['POST'])
@login_required
def admin_reset_password(user_id):
    if not current_user.can_access_admin():
        return jsonify({'success': False, 'message': 'Access denied'}), 403

    user = User.query.get_or_404(user_id)

    # Generate a secure temporary password
    alphabet = string.ascii_letters + string.digits
    temp_password = ''.join(secrets.choice(alphabet) for _ in range(10))

    try:
        user.password_hash = generate_password_hash(temp_password)
        db.session.commit()
        try:
            log_user_action(current_user.id, 'admin_reset_password', None, f"Reset password for user {user.username} ({user.id})")
        except Exception:
            pass
        # Email the user their temporary password
        subject = 'Your password has been reset'
        body = (
            f"Hello {user.get_full_name()},\n\n"
            f"An administrator has reset your account password.\n\n"
            f"Temporary password: {temp_password}\n\n"
            f"For security, please log in and change your password immediately from Settings.\n\n"
            f"If you did not expect this reset, please contact support.\n"
        )
        email_ok = False
        try:
            email_ok = send_email(subject, [user.email], body)
        except Exception as e:
            app.logger.warning(f"Password reset email send failed for {user.id}: {e}")
        msg = 'Password reset successfully.' + ('' if email_ok else ' Email delivery may have failed; share the temporary password securely.')
        return jsonify({'success': True, 'message': msg})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"admin_reset_password failed: {e}")
        return jsonify({'success': False, 'message': 'Failed to reset password'}), 500


@app.route('/admin/users/<int:user_id>/toggle_status', methods=['POST'])
@login_required
def admin_toggle_user_status(user_id):
    if not current_user.can_access_admin():
        return jsonify({'success': False, 'message': 'Access denied'}), 403

    user = User.query.get_or_404(user_id)

    # Prevent deactivating self via API
    if user.id == current_user.id:
        return jsonify({'success': False, 'message': 'You cannot change your own status.'}), 400

    try:
        user.is_active = not bool(user.is_active)
        db.session.commit()
        try:
            action = 'deactivate' if not user.is_active else 'activate'
            log_user_action(current_user.id, 'admin_toggle_user_status', None, f"{action} user {user.username} ({user.id})")
        except Exception:
            pass
        return jsonify({'success': True, 'message': ('User activated' if user.is_active else 'User deactivated'), 'is_active': user.is_active})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"admin_toggle_user_status failed: {e}")
        return jsonify({'success': False, 'message': 'Failed to change user status'}), 500


@app.route('/admin/users/<int:user_id>/delete', methods=['POST'])
@login_required
def admin_delete_user(user_id):
    if not current_user.can_access_admin():
        return jsonify({'success': False, 'message': 'Access denied'}), 403

    user = User.query.get_or_404(user_id)

    # Prevent deleting self
    if user.id == current_user.id:
        return jsonify({'success': False, 'message': 'You cannot delete your own account.'}), 400

    # Ensure at least one admin remains if deleting an admin
    try:
        admin_role = Role.query.filter_by(name='Admin').first()
        if admin_role and user.role_id == admin_role.id:
            remaining_admins = User.query.filter(User.role_id == admin_role.id, User.id != user.id).count()
            if remaining_admins <= 0:
                return jsonify({'success': False, 'message': 'Cannot delete the last admin user.'}), 400
    except Exception:
        pass

    # Block deletion if the user owns or uploaded documents
    try:
        owned_count = user.owned_documents.count()
        uploaded_count = user.uploaded_documents.count()
    except Exception:
        # Fallback safe check
        owned_count = Document.query.filter_by(owner_id=user.id).count()
        uploaded_count = Document.query.filter_by(uploaded_by_id=user.id).count()

    if owned_count > 0 or uploaded_count > 0:
        return jsonify({'success': False, 'message': f'Cannot delete user with associated documents (owned: {owned_count}, uploaded: {uploaded_count}). Reassign or remove documents first.'}), 400

    try:
        # Delete access logs referencing this user to satisfy FK constraints
        AccessLog.query.filter_by(user_id=user.id).delete(synchronize_session=False)
        db.session.delete(user)
        db.session.commit()
        try:
            log_user_action(current_user.id, 'admin_delete_user', None, f"Deleted user {user.username} ({user.id})")
        except Exception:
            pass
        return jsonify({'success': True, 'message': 'User deleted successfully.'})
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"admin_delete_user failed: {e}")
        return jsonify({'success': False, 'message': 'Failed to delete user'}), 500


@app.route('/admin/users')
@login_required
def admin_users():
    if not current_user.can_access_admin():
        flash('You do not have permission to access user management.', 'error')
        return redirect(url_for('dashboard'))
    
    page = request.args.get('page', 1, type=int)
    search = request.args.get('search', '')
    role_filter = request.args.get('role', '')
    
    query = User.query
    
    if search:
        query = query.filter(
            (User.username.contains(search)) |
            (User.email.contains(search)) |
            (User.first_name.contains(search)) |
            (User.last_name.contains(search))
        )
    
    if role_filter:
        query = query.join(Role).filter(Role.name == role_filter)
    
    users = query.order_by(User.created_at.desc()).paginate(
        page=page, per_page=20, error_out=False
    )
    
    roles = Role.query.all()
    
    return render_template('users.html',
                         users=users,
                         roles=roles,
                         search=search,
                         role_filter=role_filter)

@app.route('/admin/logs')
@login_required
def admin_activity_logs():
    if not current_user.can_access_admin():
        flash('You do not have permission to access activity logs.', 'error')
        return redirect(url_for('dashboard'))

    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 25, type=int)
    user_id = request.args.get('user_id', type=int)
    action = request.args.get('action', '', type=str)
    document_id = request.args.get('document_id', type=int)
    start = request.args.get('start', '', type=str)  # YYYY-MM-DD
    end = request.args.get('end', '', type=str)      # YYYY-MM-DD

    def _parse_date(s):
        try:
            return datetime.strptime(s, '%Y-%m-%d') if s else None
        except Exception:
            return None

    start_dt = _parse_date(start)
    end_dt = _parse_date(end)
    if end_dt:
        # make end exclusive upper bound by adding one day
        end_dt = end_dt + timedelta(days=1)

    query = AccessLog.query
    if user_id:
        query = query.filter(AccessLog.user_id == user_id)
    if action:
        query = query.filter(AccessLog.action == action)
    if document_id:
        query = query.filter(AccessLog.document_id == document_id)
    if start_dt:
        query = query.filter(AccessLog.timestamp >= start_dt)
    if end_dt:
        query = query.filter(AccessLog.timestamp < end_dt)

    logs = query.order_by(desc(AccessLog.timestamp)).paginate(page=page, per_page=per_page, error_out=False)

    # For filters
    users = User.query.order_by(User.first_name.asc(), User.last_name.asc()).all()
    actions_raw = db.session.query(AccessLog.action).distinct().all()
    actions = [a[0] for a in actions_raw if a and a[0]]

    return render_template('activity_logs.html',
                           logs=logs,
                           users=users,
                           actions=actions,
                           selected_user_id=user_id,
                           selected_action=action,
                           selected_document_id=document_id,
                           start=start,
                           end=end)

@app.route('/admin/users/<int:user_id>/activity')
@login_required
def admin_user_activity(user_id):
    if not current_user.can_access_admin():
        flash('You do not have permission to access activity logs.', 'error')
        return redirect(url_for('dashboard'))

    # Load target user
    user = User.query.get_or_404(user_id)

    # Filters and pagination
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 25, type=int)
    action = request.args.get('action', '', type=str)
    start = request.args.get('start', '', type=str)
    end = request.args.get('end', '', type=str)

    def _parse_date(s):
        try:
            return datetime.strptime(s, '%Y-%m-%d') if s else None
        except Exception:
            return None

    start_dt = _parse_date(start)
    end_dt = _parse_date(end)
    if end_dt:
        end_dt = end_dt + timedelta(days=1)

    query = AccessLog.query.filter(AccessLog.user_id == user.id)
    if action:
        query = query.filter(AccessLog.action == action)
    if start_dt:
        query = query.filter(AccessLog.timestamp >= start_dt)
    if end_dt:
        query = query.filter(AccessLog.timestamp < end_dt)

    logs = query.order_by(desc(AccessLog.timestamp)).paginate(page=page, per_page=per_page, error_out=False)

    actions_raw = db.session.query(AccessLog.action).filter(AccessLog.user_id == user.id).distinct().all()
    actions = [a[0] for a in actions_raw if a and a[0]]

    return render_template('user_activity.html',
                           target_user=user,
                           logs=logs,
                           actions=actions,
                           selected_action=action,
                           start=start,
                           end=end)

@app.route('/me/activity')
@login_required
def self_activity():
    # Show the current user's own activity log
    user = current_user

    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 25, type=int)
    action = request.args.get('action', '', type=str)
    start = request.args.get('start', '', type=str)
    end = request.args.get('end', '', type=str)

    def _parse_date(s):
        try:
            return datetime.strptime(s, '%Y-%m-%d') if s else None
        except Exception:
            return None

    start_dt = _parse_date(start)
    end_dt = _parse_date(end)
    if end_dt:
        end_dt = end_dt + timedelta(days=1)

    query = AccessLog.query.filter(AccessLog.user_id == user.id)
    if action:
        query = query.filter(AccessLog.action == action)
    if start_dt:
        query = query.filter(AccessLog.timestamp >= start_dt)
    if end_dt:
        query = query.filter(AccessLog.timestamp < end_dt)

    logs = query.order_by(desc(AccessLog.timestamp)).paginate(page=page, per_page=per_page, error_out=False)

    actions_raw = db.session.query(AccessLog.action).filter(AccessLog.user_id == user.id).distinct().all()
    actions = [a[0] for a in actions_raw if a and a[0]]

    return render_template('user_activity.html',
                           target_user=user,
                           logs=logs,
                           actions=actions,
                           selected_action=action,
                           start=start,
                           end=end)

@app.route('/settings')
@login_required
def settings():
    return render_template('settings.html')

# Error handlers
@app.errorhandler(404)
def not_found_error(error):
    return render_template('base.html', error_message="Page not found"), 404

@app.errorhandler(403)
def forbidden_error(error):
    return render_template('base.html', error_message="Access forbidden"), 403

@app.errorhandler(500)
def internal_error(error):
    db.session.rollback()
    return render_template('base.html', error_message="Internal server error"), 500
